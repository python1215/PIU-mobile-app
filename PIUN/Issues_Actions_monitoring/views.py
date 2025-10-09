from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.paginator import Paginator
from django.db.models import Q
from django.http import HttpResponse
from django.utils import timezone
from django.template.loader import render_to_string
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime

from .models import issue_action_source, IssueActions
from .forms import IssueActionSourceForm, IssueActionsForm, IssueActionsFilterForm


# ============ Issue Action Source CRUD Views ============

@login_required
def source_list(request):
    """List all issue action sources"""
    sources = issue_action_source.objects.all().order_by('-date_created')
    
    # Search functionality
    search_query = request.GET.get('search', '')
    if search_query:
        sources = sources.filter(
            Q(issue_action_source__icontains=search_query) |
            Q(loginuser__username__icontains=search_query)
        )
    
    # Pagination
    paginator = Paginator(sources, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'total_sources': sources.count()
    }
    return render(request, 'Issues_Actions_monitoring/source_list.html', context)


@login_required
def source_create(request):
    """Create new issue action source"""
    if request.method == 'POST':
        form = IssueActionSourceForm(request.POST, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Issue/Action Source created successfully!')
            return redirect('Issues_Actions_monitoring:source_list')
    else:
        form = IssueActionSourceForm(user=request.user)
    
    context = {'form': form, 'action': 'Create'}
    return render(request, 'Issues_Actions_monitoring/source_form.html', context)


@login_required
def source_update(request, pk):
    """Update existing issue action source"""
    source = get_object_or_404(issue_action_source, pk=pk)
    
    if request.method == 'POST':
        form = IssueActionSourceForm(request.POST, instance=source, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Issue/Action Source updated successfully!')
            return redirect('Issues_Actions_monitoring:source_list')
    else:
        form = IssueActionSourceForm(instance=source, user=request.user)
    
    context = {'form': form, 'action': 'Update', 'source': source}
    return render(request, 'Issues_Actions_monitoring/source_form.html', context)


@login_required
def source_delete(request, pk):
    """Delete issue action source"""
    source = get_object_or_404(issue_action_source, pk=pk)
    
    if request.method == 'POST':
        source.delete()
        messages.success(request, 'Issue/Action Source deleted successfully!')
        return redirect('Issues_Actions_monitoring:source_list')
    
    context = {'source': source}
    return render(request, 'Issues_Actions_monitoring/source_delete.html', context)


@login_required
def source_detail(request, pk):
    """View issue action source details"""
    source = get_object_or_404(issue_action_source, pk=pk)
    
    # Get related issues/actions
    related_issues = IssueActions.objects.filter(source_of_issue_or_action=source)[:5]
    
    context = {
        'source': source,
        'related_issues': related_issues,
        'total_issues': related_issues.count()
    }
    return render(request, 'Issues_Actions_monitoring/source_detail.html', context)


# ============ Issue Actions CRUD Views ============

@login_required
def issues_list(request):
    """List all issues/actions with filtering"""
    issues = IssueActions.objects.select_related(
        'project', 'year', 'quarter', 'issue_action_type', 
        'source_of_issue_or_action', 'loginUser'
    ).order_by('-date_created')
    
    # Apply filters
    filter_form = IssueActionsFilterForm(request.GET)
    if filter_form.is_valid():
        if filter_form.cleaned_data['project']:
            issues = issues.filter(project=filter_form.cleaned_data['project'])
        if filter_form.cleaned_data['year']:
            issues = issues.filter(year=filter_form.cleaned_data['year'])
        if filter_form.cleaned_data['quarter']:
            issues = issues.filter(quarter=filter_form.cleaned_data['quarter'])
        if filter_form.cleaned_data.get('issue_action_type'):
            issues = issues.filter(issue_action_type=filter_form.cleaned_data['issue_action_type'])
        if filter_form.cleaned_data['status']:
            issues = issues.filter(status=filter_form.cleaned_data['status'])
        if filter_form.cleaned_data['priority']:
            issues = issues.filter(priority=filter_form.cleaned_data['priority'])
        if filter_form.cleaned_data['assigned_to']:
            issues = issues.filter(assigned_to=filter_form.cleaned_data['assigned_to'])
        # Apply date range filters
        if filter_form.cleaned_data['assign_date_from']:
            issues = issues.filter(assign_date__gte=filter_form.cleaned_data['assign_date_from'])
        if filter_form.cleaned_data['assign_date_to']:
            issues = issues.filter(assign_date__lte=filter_form.cleaned_data['assign_date_to'])
    
    # Search functionality
    search_query = request.GET.get('search', '')
    if search_query:
        issues = issues.filter(
            Q(issue_code__icontains=search_query) |
            Q(description_of_issue_or_action__icontains=search_query) |
            Q(project__project__icontains=search_query)
        )
    
    # Pagination with configurable page size
    page_size = request.GET.get('page_size', 10)
    try:
        page_size = int(page_size)
        if page_size not in [10, 15, 25, 50, 100]:
            page_size = 10
    except (ValueError, TypeError):
        page_size = 10
    
    paginator = Paginator(issues, page_size)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'filter_form': filter_form,
        'search_query': search_query,
        'total_issues': issues.count()
    }
    return render(request, 'Issues_Actions_monitoring/issues_list.html', context)


@login_required
def issues_create(request):
    """Create new issue/action"""
    if request.method == 'POST':
        form = IssueActionsForm(request.POST, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Issue/Action created successfully!')
            return redirect('Issues_Actions_monitoring:issues_list')
    else:
        form = IssueActionsForm(user=request.user)
    
    context = {'form': form, 'action': 'Create'}
    return render(request, 'Issues_Actions_monitoring/issues_form.html', context)


@login_required
def issues_update(request, pk):
    """Update existing issue/action"""
    issue = get_object_or_404(IssueActions, pk=pk)
    
    if request.method == 'POST':
        form = IssueActionsForm(request.POST, instance=issue, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Issue/Action updated successfully!')
            return redirect('Issues_Actions_monitoring:issues_list')
    else:
        form = IssueActionsForm(instance=issue, user=request.user)
    
    context = {'form': form, 'action': 'Update', 'issue': issue}
    return render(request, 'Issues_Actions_monitoring/issues_form.html', context)


@login_required
def issues_delete(request, pk):
    """Delete issue/action"""
    issue = get_object_or_404(IssueActions, pk=pk)
    
    if request.method == 'POST':
        issue.delete()
        messages.success(request, 'Issue/Action deleted successfully!')
        return redirect('Issues_Actions_monitoring:issues_list')
    
    context = {'issue': issue}
    return render(request, 'Issues_Actions_monitoring/issues_delete.html', context)


@login_required
def issues_detail(request, pk):
    """View issue/action details"""
    issue = get_object_or_404(IssueActions, pk=pk)
    
    context = {'issue': issue}
    return render(request, 'Issues_Actions_monitoring/issues_detail.html', context)


# ============ Dashboard View ============

@login_required
def dashboard(request):
    """Issues Actions monitoring dashboard"""
    # Statistics
    total_issues = IssueActions.objects.count()
    open_issues = IssueActions.objects.filter(status='incomplete').count()
    in_progress_issues = IssueActions.objects.filter(status='in_progress').count()
    resolved_issues = IssueActions.objects.filter(status='complete').count()
    critical_issues = IssueActions.objects.filter(priority='critical').count()
    
    # Recent issues
    recent_issues = IssueActions.objects.select_related(
        'project', 'year', 'quarter', 'issue_action_type', 'source_of_issue_or_action', 'loginUser'
    ).order_by('-date_created')[:5]
    
    # Issues by priority
    priority_stats = {
        'critical': IssueActions.objects.filter(priority='critical').count(),
        'high': IssueActions.objects.filter(priority='high').count(),
        'medium': IssueActions.objects.filter(priority='medium').count(),
        'low': IssueActions.objects.filter(priority='low').count(),
    }
    
    context = {
        'total_issues': total_issues,
        'open_issues': open_issues,
        'in_progress_issues': in_progress_issues,
        'resolved_issues': resolved_issues,
        'critical_issues': critical_issues,
        'recent_issues': recent_issues,
        'priority_stats': priority_stats,
        'total_sources': issue_action_source.objects.count()
    }
    return render(request, 'Issues_Actions_monitoring/dashboard.html', context)


# ============ Export Functions ============

@login_required
def export_issues_excel(request):
    """Export issues to Excel"""
    # Get filtered queryset
    issues = IssueActions.objects.select_related(
        'project', 'year', 'quarter', 'issue_action_type',
        'source_of_issue_or_action', 'loginUser'
    ).order_by('-date_created')
    
    # Apply same filters as list view
    filter_form = IssueActionsFilterForm(request.GET)
    if filter_form.is_valid():
        if filter_form.cleaned_data['project']:
            issues = issues.filter(project=filter_form.cleaned_data['project'])
        if filter_form.cleaned_data['year']:
            issues = issues.filter(year=filter_form.cleaned_data['year'])
        if filter_form.cleaned_data['quarter']:
            issues = issues.filter(quarter=filter_form.cleaned_data['quarter'])
        if filter_form.cleaned_data.get('issue_action_type'):
            issues = issues.filter(issue_action_type=filter_form.cleaned_data['issue_action_type'])
        if filter_form.cleaned_data['status']:
            issues = issues.filter(status=filter_form.cleaned_data['status'])
        if filter_form.cleaned_data['priority']:
            issues = issues.filter(priority=filter_form.cleaned_data['priority'])
        if filter_form.cleaned_data['assigned_to']:
            issues = issues.filter(assigned_to=filter_form.cleaned_data['assigned_to'])
        # Apply date range filters
        if filter_form.cleaned_data['assign_date_from']:
            issues = issues.filter(assign_date__gte=filter_form.cleaned_data['assign_date_from'])
        if filter_form.cleaned_data['assign_date_to']:
            issues = issues.filter(assign_date__lte=filter_form.cleaned_data['assign_date_to'])
    
    # Create workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Issues Actions Report"
    
    # Define styles
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2F75B5", end_color="2F75B5", fill_type="solid")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Headers
    headers = [
        'Issue Code', 'Project', 'Year', 'Quarter', 'Issue/Action Type',
        'Description', 'Source', 'Status', 'Priority', 'Assigned To',
        'Due Date', 'Date Created', 'Date Updated', 'Resolution Notes', 'Created By'
    ]
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Data rows
    for row, issue in enumerate(issues, 2):
        data = [
            issue.issue_code,
            issue.project.project if issue.project else '',
            issue.year.profile_year if issue.year else '',
            issue.quarter.quarter if issue.quarter else '',
            issue.issue_action_type.monitoring_type if issue.issue_action_type else '',
            issue.description_of_issue_or_action,
            issue.source_of_issue_or_action.issue_action_source if issue.source_of_issue_or_action else '',
            issue.get_status_display(),
            issue.get_priority_display(),
            issue.assigned_to if issue.assigned_to else 'Unassigned',
            issue.due_date.strftime('%Y-%m-%d') if issue.due_date else '',
            issue.date_created.strftime('%Y-%m-%d %H:%M'),
            issue.date_updated.strftime('%Y-%m-%d %H:%M'),
            issue.remarks,
            issue.loginUser.username if issue.loginUser else ''
        ]
        
        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row, column=col, value=value)
            cell.border = border
            if col in [8, 9]:  # Status and Priority columns
                if issue.status == 'critical' or issue.priority == 'critical':
                    cell.fill = PatternFill(start_color="FFE6E6", end_color="FFE6E6", fill_type="solid")
                elif issue.status == 'resolved':
                    cell.fill = PatternFill(start_color="E6FFE6", end_color="E6FFE6", fill_type="solid")
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = get_column_letter(column[0].column)
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Create response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename=Issues_Actions_Report_{timestamp}.xlsx'
    
    wb.save(response)
    return response


@login_required
def export_issues_word(request):
    """Export issues to Word document"""
    # Get filtered queryset (same logic as Excel export)
    issues = IssueActions.objects.select_related(
        'project', 'year', 'quarter', 'issue_action_type',
        'source_of_issue_or_action', 'loginUser'
    ).order_by('-date_created')
    
    # Apply filters
    filter_form = IssueActionsFilterForm(request.GET)
    if filter_form.is_valid():
        if filter_form.cleaned_data['project']:
            issues = issues.filter(project=filter_form.cleaned_data['project'])
        if filter_form.cleaned_data['year']:
            issues = issues.filter(year=filter_form.cleaned_data['year'])
        if filter_form.cleaned_data['quarter']:
            issues = issues.filter(quarter=filter_form.cleaned_data['quarter'])
        if filter_form.cleaned_data.get('issue_action_type'):
            issues = issues.filter(issue_action_type=filter_form.cleaned_data['issue_action_type'])
        if filter_form.cleaned_data['status']:
            issues = issues.filter(status=filter_form.cleaned_data['status'])
        if filter_form.cleaned_data['priority']:
            issues = issues.filter(priority=filter_form.cleaned_data['priority'])
        if filter_form.cleaned_data['assigned_to']:
            issues = issues.filter(assigned_to=filter_form.cleaned_data['assigned_to'])
        # Apply date range filters
        if filter_form.cleaned_data['assign_date_from']:
            issues = issues.filter(assign_date__gte=filter_form.cleaned_data['assign_date_from'])
        if filter_form.cleaned_data['assign_date_to']:
            issues = issues.filter(assign_date__lte=filter_form.cleaned_data['assign_date_to'])
    
    # Create document
    doc = Document()
    
    # Set page margins for A4 portrait
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4 height
    section.page_width = Inches(8.27)    # A4 width
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_heading('NAWEC PIU - Issues and Actions Monitoring Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Report info
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    doc.add_paragraph(f'Total Records: {issues.count()}')
    doc.add_paragraph('')
    
    # Summary statistics
    summary = doc.add_heading('Summary Statistics', level=1)
    stats_para = doc.add_paragraph()
    stats_para.add_run(f'Open Issues: {issues.filter(status="open").count()}\n')
    stats_para.add_run(f'In Progress: {issues.filter(status="in_progress").count()}\n')
    stats_para.add_run(f'Resolved: {issues.filter(status="resolved").count()}\n')
    stats_para.add_run(f'Critical Priority: {issues.filter(priority="critical").count()}\n')
    
    # Issues details
    if issues.exists():
        details = doc.add_heading('Issues and Actions Details', level=1)
        
        for issue in issues:
            # Issue header
            issue_heading = doc.add_heading(f'Issue: {issue.issue_code}', level=2)
            
            # Create table for issue details
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Table data
            details_data = [
                ('Project', issue.project.project if issue.project else 'N/A'),
                ('Year', issue.year.profile_year if issue.year else 'N/A'),
                ('Quarter', issue.quarter.quarter if issue.quarter else 'N/A'),
                ('Type', issue.issue_action_type.monitoring_type if issue.issue_action_type else 'N/A'),
                ('Description', issue.description_of_issue_or_action),
                ('Source', issue.source_of_issue_or_action.issue_action_source if issue.source_of_issue_or_action else 'N/A'),
                ('Status', issue.get_status_display()),
                ('Priority', issue.get_priority_display()),
                ('Assigned To', issue.assigned_to if issue.assigned_to else 'Unassigned'),
                ('Assign Date', issue.assign_date.strftime('%Y-%m-%d') if issue.assign_date else 'N/A'),
                ('Due Date', issue.due_date.strftime('%Y-%m-%d') if issue.due_date else 'N/A'),
                ('Created', issue.date_created.strftime('%Y-%m-%d %H:%M')),
                ('Updated', issue.date_updated.strftime('%Y-%m-%d %H:%M')),
                ('Remarks', issue.remarks if issue.remarks else 'N/A'),
                ('Created By', issue.loginUser.username if issue.loginUser else 'N/A')
            ]
            
            for label, value in details_data:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(value)
                row_cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph('')  # Add space between issues
    
    # Save to memory buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Create response
    response = HttpResponse(
        doc_buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename=Issues_Actions_Report_{timestamp}.docx'
    
    return response