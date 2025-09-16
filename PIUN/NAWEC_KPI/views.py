from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Q, Avg, Sum, Count
from django.core.paginator import Paginator
from .models import (
    KPIIndicator, NAWEC_KPI_Monitoring, CalculateROA, CalculateNPM, CalculateMWh, 
    CalculateGAF, CalculateDSCR, CalculateTMH, CalculateATC, CalculateNECD,
    CalculateNWCD, CalculateTPS, CalculateTTP, CalculateWQCC, CalculateWQCB, CalculateNRW, CalculateDD,
    CalculateAO, CalculateDER, CalculateCR, CalculatePARI, CalculateTSQR, CalculateIMPORTS, CalculateIPP, Month
)
from .forms import KPIMonitoringDataForm, KPIIndicatorForm, CalculateROAForm, CalculateNPMForm

from setup.models import YEAR, Quarter, Indicator_Type, Measurement_Unit, Data_Collection_Frequency
from PIU_Financial_mgt.models import Project, PDO, ProjectOutCome, ProjectResult
from PIU_Mapping_project_Sites.models import projectMapping
# Using Django ORM exclusively - no SQL Server utilities needed

def calculate_achievement_gauge():
    """Calculate achievement data for the gauge visualization"""
    # Get all monitoring entries with percentage progress data
    monitoring_entries = NAWEC_KPI_Monitoring.objects.exclude(
        Percentage_progress_towards_end_target__isnull=True
    ).exclude(
        Percentage_progress_towards_end_target=0
    )
    
    if not monitoring_entries.exists():
        return {
            'achievement_rate': 0,
            'total_entries': 0,
            'performance_bands': {
                'excellent': 0,
                'good': 0,
                'average': 0,
                'weak': 0
            },
            'combined_outcomes_results': ProjectOutCome.objects.count() + ProjectResult.objects.count()
        }
    
    total_entries = monitoring_entries.count()
    
    # Calculate performance bands based on percentage progress
    excellent_count = monitoring_entries.filter(Percentage_progress_towards_end_target__gte=75).count()
    good_count = monitoring_entries.filter(
        Percentage_progress_towards_end_target__gte=56,
        Percentage_progress_towards_end_target__lt=75
    ).count()
    average_count = monitoring_entries.filter(
        Percentage_progress_towards_end_target__gte=41,
        Percentage_progress_towards_end_target__lt=56
    ).count()
    weak_count = monitoring_entries.filter(Percentage_progress_towards_end_target__lt=41).count()
    
    # Calculate average achievement rate
    average_progress = monitoring_entries.aggregate(
        avg_progress=Avg('Percentage_progress_towards_end_target')
    )['avg_progress'] or 0
    
    # Round to nearest integer for display
    achievement_rate = round(average_progress)
    
    return {
        'achievement_rate': achievement_rate,
        'total_entries': total_entries,
        'performance_bands': {
            'excellent': excellent_count,
            'good': good_count,
            'average': average_count,
            'weak': weak_count
        },
        'combined_outcomes_results': ProjectOutCome.objects.count() + ProjectResult.objects.count()
    }

@login_required
def dashboard(request):
    """NAWEC KPI Dashboard with monitoring integration"""
    # Using Django ORM exclusively
    total_monitoring_records = NAWEC_KPI_Monitoring.objects.count()
    total_indicators = KPIIndicator.objects.count()
    
    # Get current quarter entries using Django ORM
    from datetime import datetime
    current_year = datetime.now().year
    
    # Use Django ORM for queries
    current_quarter_entries = NAWEC_KPI_Monitoring.objects.filter(
        year__profile_year=current_year
    ).count()
    
    # Get project-related statistics using Django ORM
    total_pdos = PDO.objects.count()
    total_outcomes = ProjectOutCome.objects.count()
    total_results = ProjectResult.objects.count()
    
    # Calculate achievement gauge data
    achievement_data = calculate_achievement_gauge()
    
    # Recent monitoring entries for quick reference with performance calculations
    recent_entries_qs = NAWEC_KPI_Monitoring.objects.select_related(
        'project', 'year', 'quarter', 'indicator_type'
    ).order_by('-date_created')[:5]
    
    # Convert to list for template evaluation
    recent_entries = list(recent_entries_qs)
    
    # Calculate performance and variance for each entry
    for entry in recent_entries:
        if entry.End_Target_Value and entry.End_Target_Value > 0 and entry.achieved_value is not None:
            entry.performance_calculated = round((entry.achieved_value / entry.End_Target_Value) * 100, 2)
        else:
            entry.performance_calculated = None
            
        # Calculate Variance = achieved_value - End_Target_Value
        if entry.achieved_value is not None and entry.End_Target_Value is not None:
            entry.variance_calculated = round(entry.achieved_value - entry.End_Target_Value, 2)
        else:
            entry.variance_calculated = None
    
    # Recent indicators for overview - convert to list for template evaluation
    recent_indicators = list(KPIIndicator.objects.order_by('-date_created')[:6])
    
    context = {
        'total_monitoring_records': total_monitoring_records,
        'total_indicators': total_indicators,
        'current_quarter_entries': current_quarter_entries,
        'total_pdos': total_pdos,
        'total_outcomes': total_outcomes,
        'total_results': total_results,
        'achievement_data': achievement_data,
        'recent_entries': recent_entries,
        'recent_indicators': recent_indicators,
    }
    
    return render(request, 'NAWEC_KPI/dashboard.html', context)


@login_required
def kpi_management(request):
    """KPI Calculation Management dashboard"""
    
    # Get counts for each KPI calculation model
    ao_count = CalculateAO.objects.count()
    der_count = CalculateDER.objects.count()
    cr_count = CalculateCR.objects.count()
    pari_count = CalculatePARI.objects.count()
    tsqr_count = CalculateTSQR.objects.count()
    
    # Get counts for other KPI models
    roa_count = CalculateROA.objects.count()
    npm_count = CalculateNPM.objects.count()
    dd_count = CalculateDD.objects.count()
    
    total_calculations = ao_count + der_count + cr_count + pari_count + tsqr_count
    
    context = {
        'ao_count': ao_count,
        'der_count': der_count,
        'cr_count': cr_count,
        'pari_count': pari_count,
        'tsqr_count': tsqr_count,
        'roa_count': roa_count,
        'npm_count': npm_count,
        'dd_count': dd_count,
        'total_calculations': total_calculations,
    }
    
    return render(request, 'NAWEC_KPI/kpi_management.html', context)


@login_required
def performance_dashboard(request):
    """Comprehensive Performance Dashboard for NAWEC KPI Analytics"""
    # Basic statistics
    total_indicators = KPIIndicator.objects.count()
    active_monitoring = NAWEC_KPI_Monitoring.objects.count()
    
    # Calculate average performance using (achieved_value/End_Target_Value)*100
    all_entries = NAWEC_KPI_Monitoring.objects.exclude(
        achieved_value__isnull=True,
        End_Target_Value__isnull=True
    ).exclude(End_Target_Value=0)
    
    avg_performance = 0
    if all_entries.exists():
        total_performance = 0
        count = 0
        for entry in all_entries:
            if entry.End_Target_Value is not None and entry.End_Target_Value != 0:
                performance = (entry.achieved_value / entry.End_Target_Value) * 100
                total_performance += performance
                count += 1
        avg_performance = total_performance / count if count > 0 else 0
    
    # Current quarter entries
    try:
        current_year = YEAR.objects.get(profile_year=2025)  # Current year
        current_quarter = Quarter.objects.first()  # Get first quarter as default
        current_quarter_entries = NAWEC_KPI_Monitoring.objects.filter(
            year=current_year, quarter=current_quarter
        ).count()
    except:
        current_quarter_entries = 0
    
    # Overall achievement calculation
    overall_achievement = avg_performance
    
    # Determine achievement status
    if overall_achievement >= 75:
        achievement_status = 'excellent'
    elif overall_achievement >= 56:
        achievement_status = 'good'
    elif overall_achievement >= 41:
        achievement_status = 'fair'
    else:
        achievement_status = 'poor'
    
    # KPI Categories Performance (using simplified counts without complex filtering)
    total_monitoring_count = NAWEC_KPI_Monitoring.objects.count()
    quarter_split = max(1, total_monitoring_count // 4)  # Avoid division by zero
    
    kpi_categories = [
        {
            'name': 'Financial KPIs',
            'count': quarter_split,
            'avg_performance': NAWEC_KPI_Monitoring.objects.exclude(
                Percentage_progress_towards_end_target__isnull=True
            ).aggregate(avg=Avg('Percentage_progress_towards_end_target'))['avg'] or 0,
            'color': 'success',
            'trend': 'up',
            'trend_value': 2.3
        },
        {
            'name': 'Energy & Generation',
            'count': quarter_split,
            'avg_performance': NAWEC_KPI_Monitoring.objects.exclude(
                Percentage_progress_towards_end_target__isnull=True
            ).aggregate(avg=Avg('Percentage_progress_towards_end_target'))['avg'] or 0,
            'color': 'info',
            'trend': 'stable',
            'trend_value': 0.5
        },
        {
            'name': 'Service Quality',
            'count': quarter_split,
            'avg_performance': NAWEC_KPI_Monitoring.objects.exclude(
                Percentage_progress_towards_end_target__isnull=True
            ).aggregate(avg=Avg('Percentage_progress_towards_end_target'))['avg'] or 0,
            'color': 'warning',
            'trend': 'up',
            'trend_value': 1.8
        },
        {
            'name': 'Operational',
            'count': quarter_split,
            'avg_performance': NAWEC_KPI_Monitoring.objects.exclude(
                Percentage_progress_towards_end_target__isnull=True
            ).aggregate(avg=Avg('Percentage_progress_towards_end_target'))['avg'] or 0,
            'color': 'error',
            'trend': 'down',
            'trend_value': -0.3
        }
    ]
    
    # Filter parameters from GET request
    year_filter = request.GET.get('year')
    quarter_filter = request.GET.get('quarter')
    
    # Base queryset for recent entries
    entries_queryset = NAWEC_KPI_Monitoring.objects.select_related(
        'indicator_type', 'year', 'quarter', 'project'
    )
    
    # Initialize display labels for template
    selected_year_label = None
    selected_quarter_label = None
    selected_quarter_code = None
    
    # Apply year filter if provided
    if year_filter:
        try:
            # Check if it's a 4-digit year (like 2025) or a database ID
            if len(year_filter) == 4 and year_filter.isdigit():
                # Filter by profile_year for 4-digit years
                year_obj = YEAR.objects.filter(profile_year=year_filter).first()
                if year_obj:
                    entries_queryset = entries_queryset.filter(year=year_obj)
                    selected_year_label = year_obj.profile_year
            else:
                # Filter by ID for database IDs
                year_obj = YEAR.objects.filter(id=year_filter).first()
                if year_obj:
                    entries_queryset = entries_queryset.filter(year=year_obj)
                    selected_year_label = year_obj.profile_year
        except (ValueError, YEAR.DoesNotExist):
            pass
    
    # Apply quarter filter if provided
    if quarter_filter:
        try:
            # Try to get quarter by ID first
            quarter_obj = Quarter.objects.filter(id=quarter_filter).first()
            if quarter_obj:
                entries_queryset = entries_queryset.filter(quarter=quarter_obj)
                selected_quarter_label = quarter_obj.quarter
                # Extract quarter number for display (e.g., "Quarter 3" -> "Q3")
                if "Quarter" in quarter_obj.quarter:
                    quarter_num = quarter_obj.quarter.split()[-1]
                    selected_quarter_code = f"Q{quarter_num}"
            else:
                # Try to match by quarter number (1, 2, 3, 4)
                if quarter_filter in ['1', '2', '3', '4']:
                    quarter_names = {
                        '1': 'Quarter 1',
                        '2': 'Quarter 2', 
                        '3': 'Quarter 3',
                        '4': 'Quarter 4'
                    }
                    quarter_name = quarter_names.get(quarter_filter)
                    if quarter_name:
                        quarter_obj = Quarter.objects.filter(quarter=quarter_name).first()
                        if quarter_obj:
                            entries_queryset = entries_queryset.filter(quarter=quarter_obj)
                            selected_quarter_label = quarter_obj.quarter
                            # Extract quarter number for display (e.g., "Quarter 3" -> "Q3")
                            if "Quarter" in quarter_obj.quarter:
                                quarter_num = quarter_obj.quarter.split()[-1]
                                selected_quarter_code = f"Q{quarter_num}"
        except (ValueError, Quarter.DoesNotExist):
            pass
    
    
    # Get filtered entries count and recent entries with distinct to avoid duplicates
    filtered_entries_count = entries_queryset.count()
    recent_entries = list(entries_queryset.order_by('-date_created').distinct())
    
    # Calculate performance and variance for each entry with proper decimal precision
    for entry in recent_entries:
        if entry.End_Target_Value is not None and entry.End_Target_Value != 0 and entry.achieved_value is not None:
            entry.performance_calculated = round((entry.achieved_value / entry.End_Target_Value) * 100, 2)
        else:
            entry.performance_calculated = None
            
        # Calculate Variance = achieved_value - End_Target_Value
        if entry.achieved_value is not None and entry.End_Target_Value is not None:
            entry.variance_calculated = round(entry.achieved_value - entry.End_Target_Value, 2)
        else:
            entry.variance_calculated = None
            
        # Calculate Actual GIR = (Target GIR) * (Performance/100)
        # Using Targeted_Achieved_weight as Target GIR
        if (entry.Targeted_Achieved_weight is not None and 
            entry.performance_calculated is not None):
            entry.actual_gir_calculated = round((entry.Targeted_Achieved_weight * entry.performance_calculated) / 100, 2)
        else:
            entry.actual_gir_calculated = None
    
    # Calculate filtered average performance using (achieved_value/End_Target_Value)*100
    filtered_avg_performance = 0
    performance_entries = entries_queryset.exclude(
        achieved_value__isnull=True,
        End_Target_Value__isnull=True
    ).exclude(End_Target_Value=0)
    
    if performance_entries.exists():
        total_performance = 0
        count = 0
        for entry in performance_entries:
            if entry.End_Target_Value is not None and entry.End_Target_Value != 0:
                performance = (entry.achieved_value / entry.End_Target_Value) * 100
                total_performance += performance
                count += 1
        filtered_avg_performance = total_performance / count if count > 0 else 0
    
    # Calculate Overall GIR statistics
    # Overall Target GIR = sum(Target GIR) - using Targeted_Achieved_weight
    overall_target_gir = 0
    overall_actual_gir = 0
    
    target_gir_entries = entries_queryset.exclude(Targeted_Achieved_weight__isnull=True)
    for entry in target_gir_entries:
        if entry.Targeted_Achieved_weight is not None:
            overall_target_gir += entry.Targeted_Achieved_weight
    
    # Overall Actual GIR = sum(Actual GIR) - using our calculated actual_gir_calculated
    for entry in recent_entries:
        if hasattr(entry, 'actual_gir_calculated') and entry.actual_gir_calculated is not None:
            overall_actual_gir += entry.actual_gir_calculated
    
    # Overall Achievement Rate = ((Overall Actual GIR) / (Overall Target GIR)) * 100
    overall_achievement_rate = 0
    if overall_target_gir > 0:
        overall_achievement_rate = (overall_actual_gir / overall_target_gir) * 100
    
    # Get filter data for dropdowns
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    
    context = {
        'total_indicators': total_indicators,
        'active_monitoring': active_monitoring,
        'avg_performance': filtered_avg_performance,
        'current_quarter_entries': current_quarter_entries,
        'overall_achievement': overall_achievement,
        'achievement_status': achievement_status,
        'kpi_categories': kpi_categories,
        'recent_entries': recent_entries,
        'years': years,
        'quarters': quarters,
        'overall_target_gir': round(overall_target_gir, 2),
        'overall_actual_gir': round(overall_actual_gir, 2),
        'overall_achievement_rate': round(overall_achievement_rate, 2),
        'year_filter': year_filter,
        'quarter_filter': quarter_filter,
        'selected_year_label': selected_year_label,
        'selected_quarter_label': selected_quarter_label,
        'selected_quarter_code': selected_quarter_code,
        'filtered_entries_count': filtered_entries_count,
    }
    
    return render(request, 'NAWEC_KPI/performance_dashboard.html', context)

@login_required
def performance_analysis(request):
    """Analysis dashboard for KPI performance trends"""
    # Get monitoring data for analysis
    monitoring_data = NAWEC_KPI_Monitoring.objects.select_related(
        'project', 'year', 'quarter', 'indicator_type'
    ).order_by('-date_created')
    
    # Filter by year if provided
    year_filter = request.GET.get('year')
    if year_filter:
        monitoring_data = monitoring_data.filter(year__profile_year=year_filter)
    
    # Filter by project if provided
    project_filter = request.GET.get('project')
    if project_filter:
        monitoring_data = monitoring_data.filter(project__projectID=project_filter)
    
    # Paginate results
    paginator = Paginator(monitoring_data, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Get available years for filters
    available_years = YEAR.objects.all().order_by('-profile_year')
    
    context = {
        'page_obj': page_obj,
        'available_years': available_years,
        'year_filter': year_filter,
        'project_filter': project_filter,
    }
    
    return render(request, 'NAWEC_KPI/performance_analysis.html', context)

@login_required
def data_entry(request):
    """NAWEC KPI monitoring data entry form using NAWEC_KPI_Monitoring model - CREATE operation"""
    if request.method == 'POST':
        print(f'[DEBUG] POST data received: {dict(request.POST)}')
        form = KPIMonitoringDataForm(request.POST)
        
        # Set proper querysets based on POST data before validation
        project_id = request.POST.get('project')
        pdo_id = request.POST.get('pdo')
        outcome_id = request.POST.get('project_outcome')
        
        if project_id:
            form.fields['pdo'].queryset = PDO.objects.filter(project_id=project_id)
        if pdo_id:
            form.fields['project_outcome'].queryset = ProjectOutCome.objects.filter(pdo_id=pdo_id)
        if outcome_id:
            form.fields['project_result'].queryset = ProjectResult.objects.filter(project_outcome_id=outcome_id)
        
        print(f'[DEBUG] Form is_valid: {form.is_valid()}')
        if form.is_valid():
            print('[DEBUG] Form is valid, attempting to save...')
            try:
                monitoring_entry = form.save(commit=False)
                monitoring_entry.loginUser = request.user
                monitoring_entry.save()
                print(f'[DEBUG] Successfully saved monitoring entry with ID: {monitoring_entry.pk}')
                messages.success(request, 'KPI monitoring data saved successfully!')
                return redirect('NAWEC_KPI:data_entry_list')
            except Exception as save_error:
                print(f'[DEBUG] Error saving monitoring entry: {save_error}')
                messages.error(request, f'Error saving data: {str(save_error)}')
        else:
            print(f'[DEBUG] Form validation errors: {form.errors}')
            # Add form errors to messages for debugging
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = KPIMonitoringDataForm()
        
        # Initialize form querysets for proper field display
        form.fields['project'].queryset = Project.objects.filter(projectID='NAWEC')
        form.fields['pdo'].queryset = PDO.objects.none()
        form.fields['project_outcome'].queryset = ProjectOutCome.objects.none()
        form.fields['project_result'].queryset = ProjectResult.objects.none()
        form.fields['indicator_type'].queryset = Indicator_Type.objects.all()
        form.fields['indicator_description'].queryset = KPIIndicator.objects.all()
        form.fields['measurement_unit'].queryset = Measurement_Unit.objects.all()
        form.fields['collection_frequency'].queryset = Data_Collection_Frequency.objects.all()
        form.fields['year'].queryset = YEAR.objects.all()
        # Quarter field is now a ChoiceField defined in the form, not a ModelChoiceField
    
    # Get recent entries for display - convert to list for template evaluation
    recent_entries = list(NAWEC_KPI_Monitoring.objects.select_related(
        'project', 'year', 'quarter'
    ).order_by('-date_created')[:5])
    
    # Filter projects to only show NAWEC projects
    nawec_projects = Project.objects.filter(projectID__startswith='NAWEC').order_by('project')
    
    context = {
        'form': form,
        'recent_entries': recent_entries,
        'nawec_projects': nawec_projects,
    }
    
    return render(request, 'NAWEC_KPI/data_entry.html', context)

@login_required
def data_entry_list(request):
    """List all KPI monitoring entries with search and filter - READ operation"""
    entries = NAWEC_KPI_Monitoring.objects.select_related(
        'project', 'pdo', 'project_outcome', 'project_result', 'year', 'quarter', 
        'loginUser', 'indicator_type', 'measurement_unit', 
        'collection_frequency'
    ).order_by('-date_created')
    
    # Search functionality
    search_query = request.GET.get('search', '')
    if search_query:
        entries = entries.filter(
            Q(project__project__icontains=search_query) |
            Q(indicator_type__indicator_type__icontains=search_query)
        )
    
    # Filter by year
    year_filter = request.GET.get('year', '')
    if year_filter:
        entries = entries.filter(year__profile_year=year_filter)
    
    # Filter by quarter
    quarter_filter = request.GET.get('quarter', '')
    if quarter_filter:
        entries = entries.filter(quarter__id=quarter_filter)
    
    # Filter by indicator type
    indicator_type_filter = request.GET.get('indicator_type', '')
    if indicator_type_filter:
        entries = entries.filter(indicator_type__id=indicator_type_filter)
    
    # Filter by indicator description - removed field
    
    # Filter by project
    project_filter = request.GET.get('project', '')
    if project_filter:
        entries = entries.filter(project__project__icontains=project_filter)
    
    # Pagination
    paginator = Paginator(entries, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Get available filter options
    available_years = YEAR.objects.all().order_by('-profile_year')
    available_quarters = Quarter.objects.all().order_by('id')
    available_indicator_types = Indicator_Type.objects.all().order_by('indicator_type')
    # available_indicator_descriptions removed
    
    context = {
        'page_obj': page_obj,
        'available_years': available_years,
        'available_quarters': available_quarters,
        'available_indicator_types': available_indicator_types,
        # 'available_indicator_descriptions': removed,
        'search_query': search_query,
        'year_filter': year_filter,
        'quarter_filter': quarter_filter,
        'indicator_type_filter': indicator_type_filter,
        # 'indicator_description_filter': removed,
        'project_filter': project_filter,
    }
    
    return render(request, 'NAWEC_KPI/data_entry_list.html', context)

@login_required
def data_entry_export(request):
    """Export filtered KPI monitoring data to Excel"""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    from django.utils import timezone
    import io
    
    # Get the same filtered data as data_entry_list view
    entries = NAWEC_KPI_Monitoring.objects.select_related(
        'project', 'pdo', 'project_outcome', 'project_result', 'year', 'quarter', 
        'loginUser', 'indicator_type', 'measurement_unit', 
        'collection_frequency'
    ).order_by('-date_created')
    
    # Apply same filters as in data_entry_list
    search_query = request.GET.get('search', '')
    if search_query:
        entries = entries.filter(
            Q(project__project__icontains=search_query) |
            Q(indicator_type__indicator_type__icontains=search_query)
        )
    
    year_filter = request.GET.get('year', '')
    if year_filter:
        entries = entries.filter(year__profile_year=year_filter)
    
    quarter_filter = request.GET.get('quarter', '')
    if quarter_filter:
        entries = entries.filter(quarter__id=quarter_filter)
    
    indicator_type_filter = request.GET.get('indicator_type', '')
    if indicator_type_filter:
        entries = entries.filter(indicator_type__id=indicator_type_filter)
    
    # indicator_description_filter removed
    
    project_filter = request.GET.get('project', '')
    if project_filter:
        entries = entries.filter(project__project__icontains=project_filter)
    
    # Create Excel workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "KPI Data Export"
    
    # Define headers
    headers = [
        'Year', 'Quarter', 'Project', 'PDO', 'Project Outcome', 'Project Result',
        'Indicator Type', 'Indicator Description', 'Measurement Unit', 
        'Collection Frequency', 'Baseline Value', 'Achieved Value', 
        'End Target Value', 'Percentage vs Baseline', 'Percentage vs Target',
        'Targeted Weight', 'Remarks', 'Date Created', 'User'
    ]
    
    # Style for headers
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    
    # Add headers
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
    
    # Add data rows
    for row, entry in enumerate(entries, 2):
        data = [
            entry.year.profile_year if entry.year else '',
            entry.quarter.quarter if entry.quarter else '',
            entry.project.project if entry.project else '',
            entry.pdo.pdo_statement if entry.pdo else '',
            entry.project_outcome.project_outcome if entry.project_outcome else '',
            entry.project_result.project_result if entry.project_result else '',
            entry.indicator_type.indicator_type if entry.indicator_type else '',
            entry.indicator_description.indicator_description if entry.indicator_description else '',
            entry.measurement_unit.unit if entry.measurement_unit else '',
            entry.collection_frequency.frequency if entry.collection_frequency else '',
            entry.baseline_value or '',
            entry.achieved_value or '',
            entry.End_Target_Value or '',
            entry.percentage_achieved_vs_baseline or '',
            entry.percentage_achieved_vs_end_target or '',
            entry.Targeted_Achieved_weight or '',
            entry.remarks or '',
            entry.date_created.strftime('%Y-%m-%d %H:%M') if entry.date_created else '',
            entry.loginUser.username if entry.loginUser else ''
        ]
        
        for col, value in enumerate(data, 1):
            ws.cell(row=row, column=col, value=value)
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Create response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    filename = f'KPI_Data_Export_{timezone.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    # Save workbook to response
    wb.save(response)
    return response

@login_required
def data_entry_detail(request, pk):
    """View detailed KPI monitoring entry - READ operation"""
    entry = get_object_or_404(NAWEC_KPI_Monitoring, pk=pk)
    
    context = {
        'entry': entry,
    }
    
    return render(request, 'NAWEC_KPI/data_entry_detail.html', context)

@login_required
def data_entry_edit(request, pk):
    """Edit existing KPI monitoring entry - UPDATE operation"""
    entry = get_object_or_404(NAWEC_KPI_Monitoring, pk=pk)
    
    if request.method == 'POST':
        form = KPIMonitoringDataForm(request.POST, instance=entry)
        if form.is_valid():
            monitoring_entry = form.save(commit=False)
            monitoring_entry.loginUser = request.user
            monitoring_entry.save()
            messages.success(request, 'KPI monitoring data updated successfully!')
            return redirect('NAWEC_KPI:data_entry_list')
        else:
            # Add form errors to debug validation issues
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
            messages.error(request, 'Please correct the form errors and try again.')
    else:
        form = KPIMonitoringDataForm(instance=entry)
        
        # Pre-populate cascading dropdown querysets for edit mode
        if entry.project:
            form.fields['pdo'].queryset = PDO.objects.filter(project=entry.project)
            if entry.pdo:
                form.fields['project_outcome'].queryset = ProjectOutCome.objects.filter(pdo=entry.pdo)
                if entry.project_outcome:
                    form.fields['project_result'].queryset = ProjectResult.objects.filter(project_outcome=entry.project_outcome)
    
    context = {
        'form': form,
        'entry': entry,
        'is_edit': True,
    }
    
    return render(request, 'NAWEC_KPI/data_entry.html', context)

@login_required
def data_entry_delete(request, pk):
    """Delete KPI monitoring entry - DELETE operation"""
    entry = get_object_or_404(NAWEC_KPI_Monitoring, pk=pk)
    
    if request.method == 'POST':
        entry.delete()
        messages.success(request, 'KPI monitoring data deleted successfully!')
        return redirect('NAWEC_KPI:data_entry_list')
    
    context = {
        'entry': entry,
    }
    
    return render(request, 'NAWEC_KPI/data_entry_delete.html', context)

@login_required
def monitoring_list(request):
    """List all monitoring entries with filtering"""
    # Get all monitoring data with related objects
    monitoring_data = NAWEC_KPI_Monitoring.objects.select_related(
        'project', 'year', 'quarter', 'indicator_type', 'loginUser'
    ).order_by('-date_created')
    
    # Year filter
    year_filter = request.GET.get('year')
    if year_filter:
        monitoring_data = monitoring_data.filter(year__profile_year=year_filter)
    
    # Project filter
    project_filter = request.GET.get('project')
    if project_filter:
        monitoring_data = monitoring_data.filter(project__projectID=project_filter)
    
    search_query = request.GET.get('search')
    if search_query:
        monitoring_data = monitoring_data.filter(
            Q(indicator_description__icontains=search_query) |
            Q(project__project__icontains=search_query) |
            Q(remarks__icontains=search_query)
        )
    
    # Paginate results
    paginator = Paginator(monitoring_data, 25)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Get filter options
    available_years = YEAR.objects.all().order_by('-profile_year')
    
    context = {
        'page_obj': page_obj,
        'available_years': available_years,
        'year_filter': year_filter,
        'project_filter': project_filter,
        'search_query': search_query,
    }
    
    return render(request, 'NAWEC_KPI/monitoring_list.html', context)


@login_required
def indicator_create(request):
    """Create new KPI indicator"""
    if request.method == 'POST':
        form = KPIIndicatorForm(request.POST)
        if form.is_valid():
            indicator = form.save(commit=False)
            indicator.loginUser = request.user
            indicator.save()
            messages.success(request, f'KPI Indicator "{indicator.indicator_no}" created successfully!')
            return redirect('NAWEC_KPI:indicator_list')
    else:
        form = KPIIndicatorForm()
    
    context = {
        'form': form,
        'title': 'Add New KPI Indicator'
    }
    return render(request, 'NAWEC_KPI/indicator_form.html', context)


@login_required
def indicator_edit(request, pk):
    """Edit existing KPI indicator"""
    indicator = get_object_or_404(KPIIndicator, pk=pk)
    
    if request.method == 'POST':
        form = KPIIndicatorForm(request.POST, instance=indicator)
        if form.is_valid():
            form.save()
            messages.success(request, f'KPI Indicator "{indicator.indicator_no}" updated successfully!')
            return redirect('NAWEC_KPI:indicator_list')
    else:
        form = KPIIndicatorForm(instance=indicator)
    
    context = {
        'form': form,
        'indicator': indicator,
        'title': f'Edit KPI Indicator {indicator.indicator_no}'
    }
    return render(request, 'NAWEC_KPI/indicator_form.html', context)


@login_required
def indicator_list(request):
    """List all KPI indicators with search and filtering"""
    indicators = KPIIndicator.objects.select_related('loginUser').order_by('-date_created')
    
    # Search functionality
    search_query = request.GET.get('search')
    if search_query:
        indicators = indicators.filter(
            Q(indicator_no__icontains=search_query) |
            Q(indicator_description__icontains=search_query) |
            Q(attributes__icontains=search_query)
        )
    
    # Paginate results
    paginator = Paginator(indicators, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'total_indicators': indicators.count()
    }
    return render(request, 'NAWEC_KPI/indicator_list.html', context)


@login_required
def indicator_detail(request, pk):
    """View KPI indicator details with monitoring statistics"""
    indicator = get_object_or_404(KPIIndicator, pk=pk)
    
    # Get related monitoring records - convert to list for template evaluation
    monitoring_records = list(NAWEC_KPI_Monitoring.objects.filter(
        indicator_description=indicator
    ).select_related(
        'project', 'pdo', 'year', 'quarter'
    ).order_by('-date_created')[:10])
    
    # Calculate statistics
    total_monitoring = NAWEC_KPI_Monitoring.objects.filter(
        indicator_description=indicator
    ).count()
    
    # Current year and quarter statistics
    from datetime import datetime
    current_year = datetime.now().year
    current_quarter = (datetime.now().month - 1) // 3 + 1
    
    current_year_count = NAWEC_KPI_Monitoring.objects.filter(
        indicator_description=indicator,
        year__profile_year=str(current_year)
    ).count()
    
    recent_count = NAWEC_KPI_Monitoring.objects.filter(
        indicator_description=indicator,
        year__profile_year=str(current_year),
        quarter__quarter=str(current_quarter)
    ).count()
    
    context = {
        'indicator': indicator,
        'monitoring_records': monitoring_records,
        'monitoring_count': total_monitoring,
        'current_year_count': current_year_count,
        'recent_count': recent_count,
    }
    return render(request, 'NAWEC_KPI/indicator_detail.html', context)


@login_required
def indicator_delete(request, pk):
    """Delete KPI indicator"""
    indicator = get_object_or_404(KPIIndicator, pk=pk)
    
    if request.method == 'POST':
        indicator_no = indicator.indicator_no
        indicator.delete()
        messages.success(request, f'KPI Indicator "{indicator_no}" deleted successfully!')
        return redirect('NAWEC_KPI:indicator_list')
    
    context = {
        'indicator': indicator
    }
    return render(request, 'NAWEC_KPI/indicator_delete_confirm.html', context)


@login_required
def calculate_roa_popup(request):
    """Popup form for ROA calculation (KPI-01)"""
    if request.method == 'POST':
        form = CalculateROAForm(request.POST)
        if form.is_valid():
            roa_entry = form.save(commit=False)
            roa_entry.loginUser = request.user
            roa_entry.save()
            return JsonResponse({
                'success': True, 
                'message': 'ROA calculation saved successfully',
                'roa_percentage': roa_entry.roa_percentage
            })
        else:
            return JsonResponse({'success': False, 'errors': form.errors})
    else:
        form = CalculateROAForm()
    
    return render(request, 'NAWEC_KPI/calculate_roa_popup.html', {'form': form})


@login_required
def calculate_npm_popup(request):
    """Popup form for NPM calculation (KPI-02)"""
    if request.method == 'POST':
        form = CalculateNPMForm(request.POST)
        if form.is_valid():
            npm_entry = form.save(commit=False)
            npm_entry.loginUser = request.user
            npm_entry.save()
            return JsonResponse({
                'success': True, 
                'message': 'NPM calculation saved successfully',
                'net_profit_margin': npm_entry.achieved_value
            })
        else:
            return JsonResponse({'success': False, 'errors': form.errors})
    else:
        form = CalculateNPMForm()
    
    return render(request, 'NAWEC_KPI/calculate_npm_popup.html', {'form': form})


@login_required
def get_indicator_details(request, indicator_id):
    """API endpoint to get KPI indicator details including baseline, target, and weight values"""
    try:
        indicator = get_object_or_404(KPIIndicator, pk=indicator_id)
        return JsonResponse({
            'success': True,
            'indicator_no': indicator.indicator_no,
            'indicator_description': indicator.indicator_description,
            'attributes': indicator.attributes,
            'baseline_value': indicator.baseline_value,
            'End_Target_Value': indicator.End_Target_Value,
            'targeted_weight_value': indicator.targeted_weight_value
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=404)


@login_required
def get_kpi_values(request, kpi_code):
    """Get baseline and target values for a KPI based on code"""
    try:
        # Map KPI codes to their corresponding indicator numbers
        kpi_mapping = {
            'roa': 'KPI-01',
            'npm': 'KPI-02', 
            'dscr': 'KPI-03',
            'mwh': 'KPI-04',
            'gaf': 'KPI-05',
            'tde': 'KPI-06',
            'atc': 'KPI-07',
            'necd': 'KPI-08',
            'nwcd': 'KPI-09',
            'tps': 'KPI-10',
            'ttp': 'KPI-11',
            'wqcc': 'KPI-12',
            'wqcb': 'KPI-13',
            'nrw': 'KPI-14',
            'dd': 'KPI-15',
            'ao': 'KPI-16',
            'der': 'KPI-17',
            'cr': 'KPI-18',
            'pari': 'KPI-20',
            'tsqr': 'KPI-21',
        }
        
        indicator_no = kpi_mapping.get(kpi_code.lower())
        if not indicator_no:
            return JsonResponse({'success': False, 'error': 'KPI code not found'})
            
        # Find the KPI indicator by indicator number
        indicator = KPIIndicator.objects.filter(indicator_no=indicator_no).first()
        if not indicator:
            return JsonResponse({'success': False, 'error': 'KPI indicator not found in database'})
            
        return JsonResponse({
            'success': True,
            'data': {
                'baseline_value': indicator.baseline_value if indicator.baseline_value is not None else 0,
                'End_Target_Value': indicator.End_Target_Value if indicator.End_Target_Value is not None else 0,
                'indicator_no': indicator.indicator_no,
                'indicator_description': indicator.indicator_description,
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@csrf_exempt
def get_kpi_indicator_data(request, kpi_code):
    """Get KPI indicator data for form auto-population based on selected KPI code"""
    try:
        # Map KPI codes to their corresponding indicator numbers (all 20 KPIs)
        kpi_mapping = {
            'ROA': 'KPI-01',    # Return on Assets
            'NPM': 'KPI-02',    # Net Profit Margin
            'DSCR': 'KPI-03',   # Debt Service Coverage Ratio
            'MWH': 'KPI-04',    # Monthly Water Hours
            'GAF': 'KPI-05',    # Grid Availability Factor
            'TMH': 'KPI-06',    # Training Man Hours
            'ATC': 'KPI-07',    # Average Time to Connect
            'NECD': 'KPI-08',   # New Electricity Connection Delivered
            'NWCD': 'KPI-09',   # New Water Connection Delivered
            'TPS': 'KPI-10',    # Transmission and Power Supply
            'TTP': 'KPI-11',    # Tariff and Transmission Payment
            'WQCC': 'KPI-12',   # Water Quality Compliance Chemical
            'WQCB': 'KPI-13',   # Water Quality Compliance Biological
            'NRW': 'KPI-14',    # Non-Revenue Water
            'DD': 'KPI-15',     # Days Delinquent
            'AO': 'KPI-16',     # Audit Opinion
            'DER': 'KPI-17',    # Debt to Equity Ratio
            'CR': 'KPI-18',     # Current Ratio
            'PARI': 'KPI-20',   # Percentage Audit Recommendations Implementation
            'TSQR': 'KPI-21',   # Timely Submission of Quarterly Report
        }
        
        indicator_no = kpi_mapping.get(kpi_code.upper())
        if not indicator_no:
            return JsonResponse({'success': False, 'error': 'KPI code not found'})
            
        # Find the KPI indicator by indicator number
        indicator = KPIIndicator.objects.filter(indicator_no=indicator_no).first()
        
        if indicator:
            # Return actual data from database
            return JsonResponse({
                'success': True,
                'data': {
                    'baseline_value': indicator.baseline_value if indicator.baseline_value is not None else 0,
                    'End_Target_Value': indicator.End_Target_Value if indicator.End_Target_Value is not None else 0,
                    'targeted_weight_value': indicator.targeted_weight_value if indicator.targeted_weight_value is not None else 0,
                    'indicator_no': indicator.indicator_no,
                    'indicator_description': indicator.indicator_description,
                }
            })
        else:
            # Return default values for KPIs not yet in database with meaningful descriptions
            kpi_descriptions = {
                'KPI-01': 'Return on Net Assets',
                'KPI-02': 'Net Profit Margin',
                'KPI-03': 'Debt Service Coverage Ratio',
                'KPI-04': 'Monthly Water Hours',
                'KPI-05': 'Grid Availability Factor',
                'KPI-06': 'Total Debt to Equity',
                'KPI-07': 'Average Time to Connect',
                'KPI-08': 'New Electricity Connection Delivered',
                'KPI-09': 'New Water Connection Delivered',
                'KPI-10': 'Transmission and Power Supply',
                'KPI-11': 'Tariff and Transmission Payment',
                'KPI-12': 'Water Quality Compliance Chemical',
                'KPI-13': 'Water Quality Compliance Biological',
                'KPI-14': 'Non-Revenue Water',
                'KPI-15': 'Days Delinquent',
                'KPI-16': 'Audit Opinion',
                'KPI-17': 'Debt to Equity Ratio',
                'KPI-18': 'Current Ratio',
                'KPI-20': 'Percentage Audit Recommendations Implementation',
                'KPI-21': 'Timely Submission of Quarterly Report',
            }
            
            description = kpi_descriptions.get(indicator_no, f'{kpi_code} - {indicator_no}')
            
            return JsonResponse({
                'success': True,
                'data': {
                    'baseline_value': 0,
                    'End_Target_Value': 0,
                    'targeted_weight_value': 0,
                    'indicator_no': indicator_no,
                    'indicator_description': description,
                }
            })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
def get_kpi_progress_values(request, kpi_indicator_id):
    """Get KPI progress values for automatic population in data entry form"""
    try:
        kpi_indicator = KPIIndicator.objects.get(id=kpi_indicator_id)
        
        # Use original End_Target_Value from KPI Indicator only
        # DO NOT fetch compensation_end_target from old calculations to prevent memory persistence
        compensation_end_target = None
        End_Target_Value = kpi_indicator.End_Target_Value if kpi_indicator.End_Target_Value is not None else 0
        
        # Compensation end target should only be applied during active calculations
        # Not when selecting indicator for new entry
        
        return JsonResponse({
            'success': True,
            'data': {
                'baseline_value': kpi_indicator.baseline_value if kpi_indicator.baseline_value is not None else 0,
                'End_Target_Value': End_Target_Value,
                'targeted_weight_value': kpi_indicator.targeted_weight_value if kpi_indicator.targeted_weight_value is not None else 0,
                'indicator_description': kpi_indicator.indicator_description,
                'indicator_no': kpi_indicator.indicator_no,
                'compensation_end_target': compensation_end_target,
            }
        })
    except KPIIndicator.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'KPI Indicator not found'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
def calculate_roa_list(request):
    """List all ROA calculations"""
    calculations = CalculateROA.objects.select_related(
        'loginUser'
    ).order_by('-date_created')
    
    # Search functionality
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    # Paginate results
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
    }
    return render(request, 'NAWEC_KPI/calculate_roa_list.html', context)


@login_required
def calculate_npm_list(request):
    """List all NPM calculations"""
    calculations = CalculateNPM.objects.select_related(
        'loginUser'
    ).order_by('-date_created')
    
    # Search functionality
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(achieved_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    # Paginate results
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
    }
    return render(request, 'NAWEC_KPI/calculate_npm_list.html', context)


def calculate_roa_detail(request, calc_id):
    """View for displaying ROA calculation details"""
    calculation = get_object_or_404(CalculateROA, id=calc_id)
    
    context = {
        'calculation': calculation,
        'calc_type': 'ROA'
    }
    return render(request, 'NAWEC_KPI/calculation_detail.html', context)


def calculate_roa_edit(request, calc_id):
    """View for editing ROA calculations"""
    calculation = get_object_or_404(CalculateROA, id=calc_id)
    
    if request.method == 'POST':
        # Update calculation fields
        calculation.net_profit_after_tax = float(request.POST.get('net_profit_after_tax', 0))
        calculation.total_assets = float(request.POST.get('total_assets', 0))
        
        # Recalculate ROA
        if calculation.total_assets > 0:
            calculation.calculated_value = (calculation.net_profit_after_tax / calculation.total_assets) * 100
        else:
            calculation.calculated_value = 0
            
        calculation.save()
        return redirect('NAWEC_KPI:calculate_roa_list')
    
    context = {
        'calculation': calculation,
        'calc_type': 'ROA'
    }
    return render(request, 'NAWEC_KPI/calculation_edit.html', context)


def calculate_npm_detail(request, calc_id):
    """View for displaying NPM calculation details"""
    calculation = get_object_or_404(CalculateNPM, id=calc_id)
    
    context = {
        'calculation': calculation,
        'calc_type': 'NPM'
    }
    return render(request, 'NAWEC_KPI/calculation_detail.html', context)


def calculate_npm_edit(request, calc_id):
    """View for editing NPM calculations"""
    calculation = get_object_or_404(CalculateNPM, id=calc_id)
    
    if request.method == 'POST':
        # Update calculation fields
        calculation.total_revenues_turnover = float(request.POST.get('total_revenues_turnover', 0))
        calculation.netprofit = float(request.POST.get('netprofit', 0))
        calculation.compensation_amount = float(request.POST.get('compensation_amount', 0))
        
        # Recalculate NPM
        if calculation.total_revenues_turnover > 0:
            calculation.achieved_value = (calculation.netprofit / calculation.total_revenues_turnover) * 100
        else:
            calculation.achieved_value = 0
            
        calculation.save()
        return redirect('NAWEC_KPI:calculate_npm_list')
    
    context = {
        'calculation': calculation,
        'calc_type': 'NPM'
    }
    return render(request, 'NAWEC_KPI/calculation_edit.html', context)


@login_required
def get_project_outcomes(request):
    """HTMX endpoint to get Project Outcomes filtered by PDO"""
    pdo_id = request.GET.get('pdo')
    project_id = request.GET.get('project')
    
    if pdo_id:
        outcomes = ProjectOutCome.objects.filter(pdo_id=pdo_id).order_by('project_outcome')
    else:
        outcomes = ProjectOutCome.objects.none()
    
    html = '''
    <div class="mb-3">
        <label for="id_project_outcome" class="form-label">Project Outcome</label>
        <select name="project_outcome" id="id_project_outcome" class="form-control"
                hx-get="/NAWEC_KPI/get-project-results/" 
                hx-target="#project-result-field">
            <option value="">-- Select Project Outcome --</option>
    '''
    
    for outcome in outcomes:
        html += f'<option value="{outcome.pk}">{outcome.project_outcome}</option>'
    
    html += '''
        </select>
    </div>
    '''
    
    return HttpResponse(html)


@login_required
def get_project_results(request):
    """HTMX endpoint to get Project Results filtered by Project Outcome"""
    outcome_id = request.GET.get('project_outcome')
    
    if outcome_id:
        results = ProjectResult.objects.filter(project_outcome_id=outcome_id).order_by('project_result')
    else:
        results = ProjectResult.objects.none()
    
    html = '''
    <div class="mb-3">
        <label for="id_project_result" class="form-label">Project Result</label>
        <select name="project_result" id="id_project_result" class="form-control">
            <option value="">-- Select Project Result --</option>
    '''
    
    for result in results:
        html += f'<option value="{result.pk}">{result.project_result}</option>'
    
    html += '''
        </select>
    </div>
    '''
    
    return HttpResponse(html)



# DSCR CRUD Operations
@login_required
def calculate_dscr_list(request):
    """List all DSCR calculations"""
    calculations = CalculateDSCR.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(achieved_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'DSCR'
    }
    return render(request, 'NAWEC_KPI/calculate_dscr_list.html', context)

@login_required
def calculate_mwh_list(request):
    """List all MWh calculations"""
    calculations = CalculateMWh.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'MWh'
    }
    return render(request, 'NAWEC_KPI/calculate_mwh_list.html', context)

@login_required
def calculate_mwh_detail(request, calc_id):
    """View detailed MWh calculation"""
    calculation = get_object_or_404(CalculateMWh, id=calc_id)
    
    # Performance analysis
    performance_status = "excellent" if calculation.calculated_value >= 1000 else "good" if calculation.calculated_value >= 500 else "needs_improvement"
    
    context = {
        'calculation': calculation,
        'calculation_type': 'MWh',
        'performance_status': performance_status,
        'efficiency_percentage': min((calculation.calculated_value / 1000) * 100, 100) if calculation.calculated_value else 0
    }
    return render(request, 'NAWEC_KPI/calculate_mwh_detail.html', context)

@login_required
def calculate_mwh_edit(request, calc_id):
    """Edit MWh calculation"""
    calculation = get_object_or_404(CalculateMWh, id=calc_id)
    
    if request.method == 'POST':
        power_injected = float(request.POST.get('power_injected', 0))
        time_duration = float(request.POST.get('time_duration', 0))
        number_of_sources = int(request.POST.get('number_of_sources', 1))
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        
        # Calculate total energy: E_total = Σ(Ai × Bi) for i=1 to C sources
        calculated_value = power_injected * time_duration * number_of_sources
        
        # Update calculation
        calculation.power_injected = power_injected
        calculation.time_duration = time_duration
        calculation.number_of_sources = number_of_sources
        calculation.calculated_value = calculated_value
        if year_id:
            calculation.year_id = year_id
        if quarter_id:
            calculation.quarter_id = quarter_id
        calculation.save()
        
        messages.success(request, 'MWh calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_mwh_detail', calc_id=calculation.id)
    
    # Get years and quarters for dropdowns
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'MWh'
    }
    return render(request, 'NAWEC_KPI/calculate_mwh_edit.html', context)

@login_required
def calculate_gaf_list(request):
    """List all GAF calculations"""
    calculations = CalculateGAF.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'GAF'
    }
    return render(request, 'NAWEC_KPI/calculate_gaf_list.html', context)

@login_required
def calculate_gaf_detail(request, calc_id):
    """View detailed GAF calculation"""
    calculation = get_object_or_404(CalculateGAF, id=calc_id)
    
    # Performance analysis
    performance_status = "excellent" if calculation.calculated_value >= 95 else "good" if calculation.calculated_value >= 85 else "needs_improvement"
    
    context = {
        'calculation': calculation,
        'calculation_type': 'GAF',
        'performance_status': performance_status,
        'efficiency_percentage': calculation.calculated_value if calculation.calculated_value else 0
    }
    return render(request, 'NAWEC_KPI/calculate_gaf_detail.html', context)

@login_required
def calculate_gaf_edit(request, calc_id):
    """Edit GAF calculation"""
    calculation = get_object_or_404(CalculateGAF, id=calc_id)
    
    if request.method == 'POST':
        total_available_hours = float(request.POST.get('total_available_hours', 0))
        total_period_hours = float(request.POST.get('total_period_hours', 0))
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        
        # Calculate GAF: GAF = (Total Available Hours ÷ Total Period Hours) × 100
        calculated_value = (total_available_hours / total_period_hours * 100) if total_period_hours > 0 else 0
        
        # Update calculation
        calculation.total_available_hours = total_available_hours
        calculation.total_period_hours = total_period_hours
        calculation.calculated_value = calculated_value
        if year_id:
            calculation.year_id = year_id
        if quarter_id:
            calculation.quarter_id = quarter_id
        calculation.save()
        
        messages.success(request, 'GAF calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_gaf_detail', calc_id=calculation.id)
    
    # Get years and quarters for dropdowns
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'GAF'
    }
    return render(request, 'NAWEC_KPI/calculate_gaf_edit.html', context)

@login_required
def calculate_tmh_list(request):
    """List all TMH calculations"""
    calculations = CalculateTMH.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'TMH'
    }
    return render(request, 'NAWEC_KPI/calculate_tmh_list.html', context)

@login_required
def calculate_tmh_detail(request, calc_id):
    """View detailed TMH calculation"""
    calculation = get_object_or_404(CalculateTMH, id=calc_id)
    
    # Performance analysis
    performance_status = "excellent" if calculation.achieved_value >= 5 else "good" if calculation.achieved_value >= 3 else "needs_improvement"
    
    context = {
        'calculation': calculation,
        'calculation_type': 'TMH',
        'performance_status': performance_status,
        'efficiency_percentage': min((calculation.achieved_value / 5) * 100, 100) if calculation.achieved_value else 0
    }
    return render(request, 'NAWEC_KPI/calculate_tmh_detail.html', context)

@login_required
def calculate_tmh_edit(request, calc_id):
    """Edit TMH calculation with new field structure"""
    calculation = get_object_or_404(CalculateTMH, id=calc_id)
    
    if request.method == 'POST':
        title = request.POST.get('title', 'Training Session')
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        hours_per_day = float(request.POST.get('hours_per_day', 8.0))
        number_of_participants = int(request.POST.get('number_of_participants', 1))
        quarter_id = request.POST.get('quarter')
        
        # Update calculation with new fields
        calculation.title = title
        calculation.start_date = start_date
        calculation.end_date = end_date
        calculation.hours_per_day = hours_per_day
        calculation.number_of_participants = number_of_participants
        
        if quarter_id:
            calculation.quarter_id = quarter_id
            
        # The save method will automatically calculate total_man_hours using the computed properties
        calculation.save()
        
        messages.success(request, 'TMH calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_tmh_detail', calc_id=calculation.id)
    
    # Get years and quarters for dropdowns
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'TMH'
    }
    return render(request, 'NAWEC_KPI/calculate_tmh_edit.html', context)

@login_required
def calculate_atc_list(request):
    """List all ATC calculations"""
    calculations = CalculateATC.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(achieved_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'ATC'
    }
    return render(request, 'NAWEC_KPI/calculate_atc_list.html', context)

@login_required
def calculate_necd_list(request):
    """List all NECD calculations"""
    calculations = CalculateNECD.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'NECD'
    }
    return render(request, 'NAWEC_KPI/calculate_necd_list.html', context)

@login_required
def calculate_necd_detail(request, calc_id):
    """View detailed NECD calculation"""
    calculation = get_object_or_404(CalculateNECD, id=calc_id)
    
    # Performance analysis
    performance_status = "excellent" if calculation.calculated_value <= 24 else "good" if calculation.calculated_value <= 48 else "needs_improvement"
    
    context = {
        'calculation': calculation,
        'calculation_type': 'NECD',
        'performance_status': performance_status,
        'efficiency_percentage': max(100 - (calculation.calculated_value / 72) * 100, 0) if calculation.calculated_value else 100
    }
    return render(request, 'NAWEC_KPI/calculate_necd_detail.html', context)

@login_required
def calculate_necd_edit(request, calc_id):
    """Edit NECD calculation"""
    calculation = get_object_or_404(CalculateNECD, id=calc_id)
    
    if request.method == 'POST':
        total_time_taken = float(request.POST.get('total_time_taken', 0))
        total_connections = int(request.POST.get('total_connections', 0))
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        
        # Calculate NECD: NECD = Total Time Taken ÷ Total Connections
        calculated_value = (total_time_taken / total_connections) if total_connections > 0 else 0
        
        # Update calculation
        calculation.total_time_taken = total_time_taken
        calculation.total_connections = total_connections
        calculation.calculated_value = calculated_value
        if year_id:
            calculation.year_id = year_id
        if quarter_id:
            calculation.quarter_id = quarter_id
        calculation.save()
        
        messages.success(request, 'NECD calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_necd_detail', calc_id=calculation.id)
    
    # Get years and quarters for dropdowns
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'NECD'
    }
    return render(request, 'NAWEC_KPI/calculate_necd_edit.html', context)

@login_required
def calculate_nwcd_detail(request, calc_id):
    """View detailed NWCD calculation"""
    calculation = get_object_or_404(CalculateNWCD, id=calc_id)
    performance_status = "excellent" if calculation.calculated_value <= 24 else "good" if calculation.calculated_value <= 48 else "needs_improvement"
    context = {
        'calculation': calculation,
        'calculation_type': 'NWCD',
        'performance_status': performance_status,
        'efficiency_percentage': max(100 - (calculation.calculated_value / 72) * 100, 0) if calculation.calculated_value else 100
    }
    return render(request, 'NAWEC_KPI/calculate_nwcd_detail.html', context)

@login_required
def calculate_nwcd_edit(request, calc_id):
    """Edit NWCD calculation"""
    calculation = get_object_or_404(CalculateNWCD, id=calc_id)
    if request.method == 'POST':
        total_time_taken = float(request.POST.get('total_time_taken', 0))
        total_connections = int(request.POST.get('total_connections', 0))
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        calculated_value = (total_time_taken / total_connections) if total_connections > 0 else 0
        calculation.total_time_taken = total_time_taken
        calculation.total_connections = total_connections
        calculation.calculated_value = calculated_value
        if year_id:
            calculation.year_id = year_id
        if quarter_id:
            calculation.quarter_id = quarter_id
        calculation.save()
        messages.success(request, 'NWCD calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_nwcd_detail', calc_id=calculation.id)
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'NWCD'
    }
    return render(request, 'NAWEC_KPI/calculate_nwcd_edit.html', context)

@login_required
def calculate_nwcd_list(request):
    """List all NWCD calculations"""
    calculations = CalculateNWCD.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'NWCD'
    }
    return render(request, 'NAWEC_KPI/calculate_nwcd_list.html', context)

@login_required
def calculate_tps_list(request):
    """List all TPS calculations"""
    calculations = CalculateTPS.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'TPS'
    }
    return render(request, 'NAWEC_KPI/calculate_tps_list.html', context)

@login_required
def calculate_tps_detail(request, calc_id):
    """View detailed TPS calculation"""
    calculation = get_object_or_404(CalculateTPS, id=calc_id)
    performance_status = "excellent" if calculation.calculated_value >= 95 else "good" if calculation.calculated_value >= 85 else "needs_improvement"
    context = {
        'calculation': calculation,
        'calculation_type': 'TPS',
        'performance_status': performance_status,
        'efficiency_percentage': calculation.calculated_value if calculation.calculated_value else 0
    }
    return render(request, 'NAWEC_KPI/calculate_tps_detail.html', context)

@login_required
def calculate_tps_edit(request, calc_id):
    """Edit TPS calculation"""
    calculation = get_object_or_404(CalculateTPS, id=calc_id)
    if request.method == 'POST':
        number_of_on_time_payments = float(request.POST.get('number_of_on_time_payments', 0))
        total_number_of_payments_due = float(request.POST.get('total_number_of_payments_due', 0))
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        calculated_value = (number_of_on_time_payments / total_number_of_payments_due * 100) if total_number_of_payments_due > 0 else 0
        calculation.number_of_on_time_payments = number_of_on_time_payments
        calculation.total_number_of_payments_due = total_number_of_payments_due
        calculation.calculated_value = calculated_value
        if year_id:
            calculation.year_id = year_id
        if quarter_id:
            calculation.quarter_id = quarter_id
        calculation.save()
        messages.success(request, 'TPS calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_tps_detail', calc_id=calculation.id)
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'TPS'
    }
    return render(request, 'NAWEC_KPI/calculate_tps_edit.html', context)

@login_required
def calculate_ttp_list(request):
    """List all TTP calculations"""
    calculations = CalculateTTP.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'TTP'
    }
    return render(request, 'NAWEC_KPI/calculate_ttp_list.html', context)

@login_required
def calculate_ttp_detail(request, calc_id):
    """View detailed TTP calculation"""
    calculation = get_object_or_404(CalculateTTP, id=calc_id)
    performance_status = "excellent" if calculation.calculated_value >= 95 else "good" if calculation.calculated_value >= 85 else "needs_improvement"
    context = {
        'calculation': calculation,
        'calculation_type': 'TTP',
        'performance_status': performance_status,
        'efficiency_percentage': calculation.calculated_value if calculation.calculated_value else 0
    }
    return render(request, 'NAWEC_KPI/calculate_ttp_detail.html', context)

@login_required
def calculate_ttp_edit(request, calc_id):
    """Edit TTP calculation"""
    calculation = get_object_or_404(CalculateTTP, id=calc_id)
    if request.method == 'POST':
        timely_payments = int(request.POST.get('timely_payments', 0))
        total_payments_due = int(request.POST.get('total_payments_due', 0))
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        calculated_value = (timely_payments / total_payments_due) if total_payments_due > 0 else 0
        calculation.timely_payments = timely_payments
        calculation.total_payments_due = total_payments_due
        calculation.calculated_value = calculated_value
        if year_id:
            calculation.year_id = year_id
        if quarter_id:
            calculation.quarter_id = quarter_id
        calculation.save()
        messages.success(request, 'TTP calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_ttp_detail', calc_id=calculation.id)
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'TTP'
    }
    return render(request, 'NAWEC_KPI/calculate_ttp_edit.html', context)

@login_required
def calculate_wqcc_list(request):
    """List all WQCC calculations"""
    calculations = CalculateWQCC.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'WQCC'
    }
    return render(request, 'NAWEC_KPI/calculate_wqcc_list.html', context)

@login_required
def calculate_wqcc_detail(request, calc_id):
    """View detailed WQCC calculation"""
    calculation = get_object_or_404(CalculateWQCC, id=calc_id)
    performance_status = "excellent" if calculation.calculated_value >= 95 else "good" if calculation.calculated_value >= 85 else "needs_improvement"
    context = {
        'calculation': calculation,
        'calculation_type': 'WQCC',
        'performance_status': performance_status,
        'efficiency_percentage': calculation.calculated_value if calculation.calculated_value else 0
    }
    return render(request, 'NAWEC_KPI/calculate_wqcc_detail.html', context)

@login_required
def calculate_wqcc_edit(request, calc_id):
    """Edit WQCC calculation"""
    calculation = get_object_or_404(CalculateWQCC, id=calc_id)
    if request.method == 'POST':
        compliant_samples = int(request.POST.get('number_of_compliant_water_samples', 0))
        total_samples = int(request.POST.get('total_number_of_tested_water_samples', 0))
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        calculated_value = (compliant_samples / total_samples * 100) if total_samples > 0 else 0
        calculation.number_of_compliant_water_samples = compliant_samples
        calculation.total_number_of_tested_water_samples = total_samples
        calculation.calculated_value = calculated_value
        if year_id:
            calculation.year_id = year_id
        if quarter_id:
            calculation.quarter_id = quarter_id
        calculation.save()
        messages.success(request, 'WQCC calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_wqcc_detail', calc_id=calculation.id)
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'WQCC'
    }
    return render(request, 'NAWEC_KPI/calculate_wqcc_edit.html', context)

@login_required
def calculate_wqcb_list(request):
    """List all WQCB calculations"""
    calculations = CalculateWQCB.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'WQCB'
    }
    return render(request, 'NAWEC_KPI/calculate_wqcb_list.html', context)

@login_required
def calculate_wqcb_detail(request, calc_id):
    """View detailed WQCB calculation"""
    calculation = get_object_or_404(CalculateWQCB, id=calc_id)
    performance_status = "excellent" if calculation.calculated_value >= 95 else "good" if calculation.calculated_value >= 85 else "needs_improvement"
    context = {
        'calculation': calculation,
        'calculation_type': 'WQCB',
        'performance_status': performance_status,
        'efficiency_percentage': calculation.calculated_value if calculation.calculated_value else 0
    }
    return render(request, 'NAWEC_KPI/calculate_wqcb_detail.html', context)

@login_required
def calculate_wqcb_edit(request, calc_id):
    """Edit WQCB calculation"""
    calculation = get_object_or_404(CalculateWQCB, id=calc_id)
    if request.method == 'POST':
        compliant_samples = int(request.POST.get('number_of_compliant_water_samples', 0))
        total_samples = int(request.POST.get('total_number_of_tested_water_samples', 0))
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        calculated_value = (compliant_samples / total_samples * 100) if total_samples > 0 else 0
        calculation.number_of_compliant_water_samples = compliant_samples
        calculation.total_number_of_tested_water_samples = total_samples
        calculation.calculated_value = calculated_value
        if year_id:
            calculation.year_id = year_id
        if quarter_id:
            calculation.quarter_id = quarter_id
        calculation.save()
        messages.success(request, 'WQCB calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_wqcb_detail', calc_id=calculation.id)
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'WQCB'
    }
    return render(request, 'NAWEC_KPI/calculate_wqcb_edit.html', context)

@login_required
def calculate_nrw_list(request):
    """List all NRW calculations"""
    calculations = CalculateNRW.objects.select_related('loginUser', 'year', 'quarter').order_by('-date_created')
    
    search_query = request.GET.get('search')
    if search_query:
        calculations = calculations.filter(
            Q(calculated_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    paginator = Paginator(calculations, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'total_calculations': calculations.count(),
        'search_query': search_query,
        'calculation_type': 'NRW'
    }
    return render(request, 'NAWEC_KPI/calculate_nrw_list.html', context)

@login_required
def calculate_nrw_detail(request, calc_id):
    """View detailed NRW calculation"""
    calculation = get_object_or_404(CalculateNRW, id=calc_id)
    performance_status = "excellent" if calculation.calculated_value <= 15 else "good" if calculation.calculated_value <= 25 else "needs_improvement"
    context = {
        'calculation': calculation,
        'calculation_type': 'NRW',
        'performance_status': performance_status,
        'efficiency_percentage': max(100 - calculation.calculated_value, 0) if calculation.calculated_value else 100
    }
    return render(request, 'NAWEC_KPI/calculate_nrw_detail.html', context)

@login_required
def calculate_nrw_edit(request, calc_id):
    """Edit NRW calculation"""
    calculation = get_object_or_404(CalculateNRW, id=calc_id)
    if request.method == 'POST':
        water_entering_system = float(request.POST.get('water_entering_system', 0))
        billed_authorized_consumption = float(request.POST.get('billed_authorized_consumption', 0))
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        calculated_value = (water_entering_system / billed_authorized_consumption * 100) if billed_authorized_consumption > 0 else 0
        calculation.water_entering_system = water_entering_system
        calculation.billed_authorized_consumption = billed_authorized_consumption
        calculation.calculated_value = calculated_value
        if year_id:
            calculation.year_id = year_id
        if quarter_id:
            calculation.quarter_id = quarter_id
        calculation.save()
        messages.success(request, 'NRW calculation updated successfully!')
        return redirect('NAWEC_KPI:calculate_nrw_detail', calc_id=calculation.id)
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.all().order_by('quarter')
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters,
        'calculation_type': 'NRW'
    }
    return render(request, 'NAWEC_KPI/calculate_nrw_edit.html', context)

# ATC Edit View
def calculate_atc_edit(request, calc_id):
    """View for editing ATC calculations"""
    calculation = get_object_or_404(CalculateATC, id=calc_id, loginUser=request.user)
    
    if request.method == 'POST':
        # Update calculation fields
        calculation.billing_efficiency = float(request.POST.get('billing_efficiency', 0))
        calculation.collection_efficiency = float(request.POST.get('collection_efficiency', 0))
        
        # Recalculate ATC
        calculation.calculated_value = (1 - (calculation.billing_efficiency * calculation.collection_efficiency / 10000)) * 100
        
        calculation.save()
        return redirect('NAWEC_KPI:calculate_atc_list')
    
    context = {
        'calculation': calculation,
        'calculation_type': 'ATC'
    }
    return render(request, 'NAWEC_KPI/calculate_atc_edit.html', context)

# ATC Detail View
def calculate_atc_detail(request, calc_id):
    """View for showing ATC calculation details"""
    calculation = get_object_or_404(CalculateATC, id=calc_id)
    
    context = {
        'calculation': calculation,
        'calculation_type': 'ATC'
    }
    return render(request, 'NAWEC_KPI/calculate_atc_detail.html', context)

# DSCR Edit View
def calculate_dscr_edit(request, calc_id):
    """View for editing DSCR calculations"""
    calculation = get_object_or_404(CalculateDSCR, unique_id=calc_id, loginUser=request.user)
    
    if request.method == 'POST':
        # Update calculation fields
        calculation.debt_service = float(request.POST.get('debt_service', 0))
        calculation.cashflow = float(request.POST.get('cashflow', 0))
        
        # Recalculate DSCR
        if calculation.debt_service > 0:
            calculation.calculated_value = calculation.cashflow / calculation.debt_service
        else:
            calculation.calculated_value = 0
            
        calculation.save()
        return redirect('NAWEC_KPI:calculate_dscr_list')
    
    context = {
        'calculation': calculation,
        'calculation_type': 'DSCR'
    }
    return render(request, 'NAWEC_KPI/calculate_dscr_edit.html', context)

# DSCR Detail View
def calculate_dscr_detail(request, calc_id):
    """View for showing DSCR calculation details"""
    calculation = get_object_or_404(CalculateDSCR, unique_id=calc_id)
    
    context = {
        'calculation': calculation,
        'calculation_type': 'DSCR'
    }
    return render(request, 'NAWEC_KPI/calculate_dscr_detail.html', context)


@login_required
def performance_report(request):
    """Generate comprehensive performance report with all KPI data and calculations"""
    from datetime import datetime
    import io
    from django.http import HttpResponse
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    
    # Get filter parameters
    year_filter = request.GET.get('year')
    quarter_filter = request.GET.get('quarter')
    format_type = request.GET.get('format', 'html')  # Default to HTML report
    
    # Base queryset for monitoring entries - show all KPI data for comprehensive reporting
    entries_queryset = NAWEC_KPI_Monitoring.objects.select_related(
        'project', 'pdo', 'project_outcome', 'project_result', 
        'indicator_type', 'year', 'quarter', 'loginUser'
    )
    
    # Apply filters if provided
    if year_filter:
        entries_queryset = entries_queryset.filter(year_id=year_filter)
    if quarter_filter:
        # Handle numeric quarter values (1, 2, 3, 4) with correct mapping
        if quarter_filter in ['1', '2', '3', '4']:
            # Map numeric quarters to actual Quarter database IDs
            quarter_mapping = {
                '1': 10030,  # Quarter 1 (ID: 10030)
                '2': 10031,  # Quarter 2 (ID: 10031)
                '3': 10032,  # Quarter 3 (ID: 10032)
                '4': 10033   # Quarter 4 (ID: 10033)
            }
            quarter_id = quarter_mapping.get(quarter_filter)
            if quarter_id:
                try:
                    quarter_obj = Quarter.objects.get(id=quarter_id)
                    entries_queryset = entries_queryset.filter(quarter=quarter_obj)
                except Quarter.DoesNotExist:
                    pass
        else:
            # Fallback for direct ID filtering
            entries_queryset = entries_queryset.filter(quarter_id=quarter_filter)
    
    # Get all KPI indicators to ensure comprehensive reporting
    all_kpi_indicators = KPIIndicator.objects.all().order_by('indicator_no')
    
    # Get filtered data with calculations and create comprehensive KPI report structure
    report_entries = []
    comprehensive_kpi_data = {}
    overall_target_gir = 0
    overall_actual_gir = 0
    
    # Initialize comprehensive structure with all KPIs
    for kpi_indicator in all_kpi_indicators:
        indicator_name = kpi_indicator.indicator_description
        comprehensive_kpi_data[indicator_name] = {
            'indicator_no': kpi_indicator.indicator_no,
            'baseline_value': kpi_indicator.baseline_value,
            'target_value': kpi_indicator.End_Target_Value,
            'quarters': {}
        }
    
    # Process monitoring entries and organize by indicator and quarter
    for entry in entries_queryset:
        # Calculate Performance, Variance, and Actual GIR
        performance_calculated = None
        variance_calculated = None
        actual_gir_calculated = None
        
        if entry.achieved_value is not None and entry.End_Target_Value is not None and entry.End_Target_Value != 0:
            performance_calculated = round((entry.achieved_value / entry.End_Target_Value) * 100, 2)
            variance_calculated = round(entry.achieved_value - entry.End_Target_Value, 2)
            
            if entry.Targeted_Achieved_weight is not None:
                actual_gir_calculated = round((entry.Targeted_Achieved_weight * performance_calculated) / 100, 2)
                overall_actual_gir += actual_gir_calculated
        
        if entry.Targeted_Achieved_weight is not None:
            overall_target_gir += entry.Targeted_Achieved_weight
        
        # Add calculated values to entry
        entry.performance_calculated = performance_calculated
        entry.variance_calculated = variance_calculated
        entry.actual_gir_calculated = actual_gir_calculated
        
        report_entries.append(entry)
        
        # Group entries by indicator and quarter for comprehensive display
        if hasattr(entry, 'indicator_description') and entry.indicator_description:
            indicator_name = entry.indicator_description.indicator_description
            quarter_key = f"Q{entry.quarter.id}"
            
            # Map quarter IDs to standard Q1, Q2, Q3, Q4 format
            quarter_mapping = {
                10030: 'Q1',  # Quarter 1
                10031: 'Q2',  # Quarter 2  
                10032: 'Q3',  # Quarter 3
                10033: 'Q4'   # Quarter 4
            }
            quarter_display = quarter_mapping.get(entry.quarter.id, f"Q{entry.quarter.id}")
            
            if indicator_name in comprehensive_kpi_data:
                comprehensive_kpi_data[indicator_name]['quarters'][quarter_display] = {
                    'baseline_value': entry.baseline_value,
                    'target_value': entry.End_Target_Value,
                    'achieved_value': entry.achieved_value,
                    'performance_calculated': performance_calculated,
                    'variance_calculated': variance_calculated,
                    'status': entry
                }
    
    # Calculate Overall Achievement Rate
    overall_achievement_rate = 0
    if overall_target_gir > 0:
        overall_achievement_rate = round((overall_actual_gir / overall_target_gir) * 100, 2)
    
    # Check format type and render HTML report if requested
    if format_type == 'html':
        # Calculate needle angle for gauge (180 degrees for semicircle)
        # 0% = -90 degrees, 100% = 90 degrees
        needle_angle = (overall_achievement_rate * 1.8) - 90
        
        # Prepare filter information string
        filter_info = ""
        if year_filter:
            try:
                year_obj = YEAR.objects.get(id=year_filter)
                filter_info += f"Year: {year_obj.profile_year} "
            except:
                filter_info += f"Year: {year_filter} "
        if quarter_filter:
            try:
                quarter_obj = Quarter.objects.get(id=quarter_filter)
                filter_info += f"Quarter: Q{quarter_obj.quarter} "
            except:
                filter_info += f"Quarter: {quarter_filter} "
        if not filter_info:
            filter_info = "No Filters Applied"
        
        # Prepare context for HTML template
        context = {
            'report_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'filter_info': filter_info.strip(),
            'overall_target_gir': overall_target_gir,
            'overall_actual_gir': overall_actual_gir,
            'overall_achievement_rate': overall_achievement_rate,
            'needle_angle': needle_angle,
            'report_entries': report_entries,
            'comprehensive_kpi_data': comprehensive_kpi_data,
        }
        
        return render(request, 'NAWEC_KPI/performance_report.html', context)
    
    # Handle Word format
    elif format_type == 'word':
        from docx import Document
        from docx.shared import Inches, Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import RGBColor
        
        # Create document
        doc = Document()
        
        # Set document margins for A4 portrait
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Document title
        title = doc.add_heading('NAWEC KPI PERFORMANCE REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Report metadata
        metadata_p = doc.add_paragraph()
        metadata_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata_p.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        
        # Filter info
        filter_info = ""
        if year_filter:
            try:
                year_obj = YEAR.objects.get(id=year_filter)
                filter_info += f"Year: {year_obj.profile_year} "
            except:
                filter_info += f"Year: {year_filter} "
        if quarter_filter:
            try:
                quarter_obj = Quarter.objects.get(id=quarter_filter)
                filter_info += f"Quarter: Q{quarter_obj.quarter} "
            except:
                filter_info += f"Quarter: {quarter_filter} "
        if not filter_info:
            filter_info = "No Filters Applied"
        
        metadata_p.add_run(f'Filters: {filter_info.strip()}')
        
        # Overall Performance Summary
        doc.add_heading('Overall Performance Summary', level=1)
        
        # Performance metrics
        metrics_p = doc.add_paragraph()
        metrics_p.add_run('Overall Target GIR: ').bold = True
        metrics_p.add_run(f'{overall_target_gir:.2f}\n')
        metrics_p.add_run('Overall Actual GIR: ').bold = True
        metrics_p.add_run(f'{overall_actual_gir:.2f}\n')
        metrics_p.add_run('Overall Achievement Rate: ').bold = True
        achievement_run = metrics_p.add_run(f'{overall_achievement_rate:.2f}%')
        achievement_run.bold = True
        
        # Color code achievement rate
        if overall_achievement_rate >= 75:
            achievement_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif overall_achievement_rate >= 50:
            achievement_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            achievement_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Achievement Gauge Visual - Circular Representation
        gauge_section = doc.add_paragraph()
        gauge_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gauge_section.add_run('Achievement Gauge Visual\n').bold = True
        
        def create_dashboard_gauge(rate):
            """Create a visual gauge representation for Word documents"""
            # Determine performance band and color
            if rate >= 75:
                band = "Excellent"
                color = "🟢"
            elif rate >= 56:
                band = "Good"  
                color = "🟡"
            elif rate >= 41:
                band = "Average"
                color = "🟠"
            else:
                band = "Weak"
                color = "🔴"
            
            # Create segments for the gauge (16 segments total)
            total_segments = 16
            filled_segments = int((rate / 100) * total_segments)
            
            segments = []
            for i in range(total_segments):
                if i < filled_segments:
                    segments.append("●")  # Filled segment
                else:
                    segments.append("○")  # Empty segment
            
            # Create the gauge visual
            gauge_visual = f"""
    ╭─────────────────────────╮
    │        {segments[0]}{segments[1]}{segments[2]}{segments[3]}        │
    │      {segments[15]}       {segments[4]}      │
    │    {segments[14]}           {segments[5]}    │
    │  {segments[13]}               {segments[6]}  │
    │ {segments[12]}                 {segments[7]} │
    │ {segments[11]}       ●         {segments[8]} │
    │ {segments[10]}                 {segments[9]} │
    │                         │
    │        ACHIEVEMENT      │
    │ {segments[15]}     {rate:.1f}%     {segments[8]} │
    │ {segments[14]}   {band}   {segments[9]} │
    │   {segments[13]}{segments[12]}{segments[11]}{segments[10]}    {color}     │
    │                         │
    ╰─────────────────────────╯
    
Performance Bands:
🟢 Excellent: 75-100%  🟡 Good: 56-74%
🟠 Average: 41-55%     🔴 Weak: 0-40%

Gauge shows {filled_segments} of {total_segments} segments filled
({rate:.1f}% achievement rate)
"""
            return gauge_visual
        
        # Add the dashboard-style gauge
        gauge_para = doc.add_paragraph()
        gauge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gauge_run = gauge_para.add_run(create_dashboard_gauge(overall_achievement_rate))
        gauge_run.font.name = 'Courier New'  # Monospace for alignment
        gauge_run.font.size = Pt(10)
        
        # Achievement Gauge Description
        gauge_desc = doc.add_paragraph()
        gauge_desc.add_run('Performance Analysis:\n').bold = True
        if overall_achievement_rate >= 75:
            gauge_desc.add_run('🟢 Excellent Performance (75-100%): Outstanding achievement of targets.')
        elif overall_achievement_rate >= 50:
            gauge_desc.add_run('🟡 Good Performance (50-74%): Satisfactory progress towards targets.')
        elif overall_achievement_rate >= 25:
            gauge_desc.add_run('🟠 Fair Performance (25-49%): Moderate progress, improvement needed.')
        else:
            gauge_desc.add_run('🔴 Weak Performance (0-24%): Significant improvement required.')
        
        # Performance Bands Legend
        legend_p = doc.add_paragraph()
        legend_p.add_run('Performance Bands:\n').bold = True
        legend_p.add_run('🟢 Excellent: 75-100% | 🟡 Good: 56-74% | 🟠 Average: 41-55% | 🔴 Weak: 0-40%')
        
        # Gauge Interpretation
        interpretation_p = doc.add_paragraph()
        interpretation_p.add_run('Gauge Interpretation:\n').bold = True
        interpretation_p.add_run(f'The gauge shows {overall_achievement_rate:.1f}% achievement rate. ')
        interpretation_p.add_run(f'Each filled segment (●) represents progress. ')
        interpretation_p.add_run(f'Empty segments (○) show remaining progress needed to reach 100%.')
        
        # Gauge Interpretation
        interpretation_p = doc.add_paragraph()
        interpretation_p.add_run('Gauge Interpretation:\n').bold = True
        interpretation_p.add_run(f'The gauge shows {overall_achievement_rate:.1f}% achievement rate. ')
        interpretation_p.add_run(f'Each filled segment (█) represents 5% progress. ')
        interpretation_p.add_run(f'Empty segments (░) show remaining progress needed to reach 100%.')
        
        # KPI Data Table
        if report_entries:
            doc.add_heading('Detailed KPI Performance Data', level=1)
            
            # Create table with compact columns for A4 portrait
            table = doc.add_table(rows=1, cols=10)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header row
            hdr_cells = table.rows[0].cells
            headers = ['Indicator', 'Project', 'PDO', 'Base', 'Target', 'Achieved', 'Perf%', 'Var', 'T.Wgt', 'A.Wgt']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows
            for entry in report_entries:
                row_cells = table.add_row().cells
                
                # Truncate long text for better fit
                indicator_text = str(entry.indicator_description)[:15] + "..." if len(str(entry.indicator_description)) > 15 else str(entry.indicator_description)
                project_text = str(entry.project)[:10] + "..." if len(str(entry.project)) > 10 else str(entry.project)
                pdo_text = str(entry.pdo)[:10] + "..." if len(str(entry.pdo)) > 10 else str(entry.pdo)
                
                row_data = [
                    indicator_text,
                    project_text,
                    pdo_text,
                    f'{entry.baseline_value:.1f}' if entry.baseline_value else 'N/A',
                    f'{entry.End_Target_Value:.1f}' if entry.End_Target_Value else 'N/A',
                    f'{entry.achieved_value:.1f}' if entry.achieved_value else 'N/A',
                    f'{entry.performance_calculated:.0f}%' if entry.performance_calculated else 'N/A',
                    f'{entry.variance_calculated:.1f}' if entry.variance_calculated else 'N/A',
                    f'{entry.Targeted_Achieved_weight:.1f}' if entry.Targeted_Achieved_weight else 'N/A',
                ]
                
                for i, data in enumerate(row_data):
                    row_cells[i].text = str(data)
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document to memory
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Create response
        response = HttpResponse(
            doc_io.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filter_suffix = ""
        if year_filter or quarter_filter:
            filter_suffix = f"_{year_filter if year_filter else 'AllYears'}_{quarter_filter if quarter_filter else 'AllQuarters'}"
        
        filename = f'NAWEC_KPI_Performance_Report{filter_suffix}_{timestamp}.docx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    # Create Excel workbook (for Excel format)
    wb = Workbook()
    ws = wb.active
    ws.title = "NAWEC KPI Performance Report"
    
    # Define styles
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    subheader_font = Font(bold=True, color="000000")
    subheader_fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
    
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Report header
    ws.merge_cells('A1:R3')
    ws['A1'] = "NAWEC KPI PERFORMANCE REPORT"
    ws['A1'].font = Font(bold=True, size=16, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
    
    # Report metadata
    current_row = 5
    ws[f'A{current_row}'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ws[f'A{current_row}'].font = Font(bold=True)
    
    # Filter information
    current_row += 1
    filter_info = "Filters Applied: "
    if year_filter:
        try:
            year_obj = YEAR.objects.get(id=year_filter)
            filter_info += f"Year: {year_obj.profile_year} "
        except:
            filter_info += f"Year: {year_filter} "
    if quarter_filter:
        try:
            quarter_obj = Quarter.objects.get(id=quarter_filter)
            filter_info += f"Quarter: Q{quarter_obj.quarter} "
        except:
            filter_info += f"Quarter: {quarter_filter} "
    if not year_filter and not quarter_filter:
        filter_info += "None (All Data)"
    
    ws[f'A{current_row}'] = filter_info
    ws[f'A{current_row}'].font = Font(italic=True)
    
    # Overall GIR Summary
    current_row += 2
    ws[f'A{current_row}'] = "OVERALL GIR SUMMARY"
    ws[f'A{current_row}'].font = subheader_font
    ws[f'A{current_row}'].fill = subheader_fill
    
    current_row += 1
    ws[f'A{current_row}'] = f"Overall Target GIR: {round(overall_target_gir, 2)}"
    ws[f'A{current_row}'].font = Font(bold=True)
    
    current_row += 1
    ws[f'A{current_row}'] = f"Overall Actual GIR: {round(overall_actual_gir, 2)}"
    ws[f'A{current_row}'].font = Font(bold=True)
    
    current_row += 1
    ws[f'A{current_row}'] = f"Overall Achievement Rate: {overall_achievement_rate}%"
    ws[f'A{current_row}'].font = Font(bold=True, color="0070C0")
    
    # Data table headers
    current_row += 3
    headers = [
        'Project', 'PDO', 'Outcome', 'Result', 'Indicator Type', 'Description',
        'Baseline Value', 'Achieved Value', 'End Target Value', 'Performance (%)',
        'Variance', 'Target GIR', 'Actual GIR', 'Achieved Weight',
        'Progress from Baseline (%)', 'Progress to Target (%)', 'Year', 'Quarter'
    ]
    
    # Write headers
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=current_row, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = border
    
    # Write data rows
    for entry in report_entries:
        current_row += 1
        data_row = [
            entry.project.project if entry.project else '',
            entry.pdo.pdo_statement if entry.pdo else '',
            entry.project_outcome.project_outcome if entry.project_outcome else '',
            entry.project_result.project_result if entry.project_result else '',
            entry.indicator_type.indicator_type if entry.indicator_type else '',
            entry.indicator_description or '',
            entry.baseline_value or '',
            entry.achieved_value or '',
            entry.End_Target_Value or '',
            entry.performance_calculated or '',
            entry.variance_calculated or '',
            entry.Targeted_Achieved_weight or '',
            entry.actual_gir_calculated or '',
            entry.Percentage_progress_from_baseline or '',
            entry.Percentage_progress_towards_end_target or '',
            entry.year.profile_year if entry.year else '',
            f"Q{entry.quarter.quarter}" if entry.quarter else ''
        ]
        
        for col, value in enumerate(data_row, 1):
            # Convert all values to string or number for Excel compatibility
            if value is None:
                excel_value = ''
            elif isinstance(value, (int, float)):
                excel_value = value
            else:
                excel_value = str(value)
            
            cell = ws.cell(row=current_row, column=col, value=excel_value)
            cell.border = border
            
            # Apply conditional formatting for variance
            if col == 11 and isinstance(value, (int, float)):  # Variance column
                if value >= 0:
                    cell.font = Font(color="008000")  # Green for positive
                else:
                    cell.font = Font(color="FF0000")  # Red for negative
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = get_column_letter(column[0].column)
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Generate filename with timestamp and filters
    filename_parts = ["NAWEC_KPI_Performance_Report"]
    if year_filter:
        try:
            year_obj = YEAR.objects.get(id=year_filter)
            filename_parts.append(f"Year_{year_obj.profile_year}")
        except:
            filename_parts.append(f"Year_{year_filter}")
    if quarter_filter:
        try:
            quarter_obj = Quarter.objects.get(id=quarter_filter)
            filename_parts.append(f"Q{quarter_obj.quarter}")
        except:
            filename_parts.append(f"Q{quarter_filter}")
    
    filename_parts.append(datetime.now().strftime("%Y%m%d_%H%M%S"))
    filename = "_".join(filename_parts) + ".xlsx"
    
    # Create HTTP response with Excel file
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    # Save workbook to response
    wb.save(response)
    return response


# ======================= DD (Debtor Days) CRUD Views =======================

@login_required
def calculate_dd_list(request):
    """List all DD calculations with search and pagination"""
    search_query = request.GET.get('search', '')
    
    calculations = CalculateDD.objects.all().order_by('-date_created')
    
    if search_query:
        calculations = calculations.filter(
            Q(achieved_value__icontains=search_query) |
            Q(year__profile_year__icontains=search_query) |
            Q(quarter__quarter__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    # Pagination
    paginator = Paginator(calculations, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'total_count': calculations.count()
    }
    return render(request, 'NAWEC_KPI/calculate_dd_list.html', context)


@login_required
def calculate_dd_detail(request, calc_id):
    """View DD calculation details"""
    calculation = get_object_or_404(CalculateDD, pk=calc_id)
    
    # Performance analysis
    if calculation.achieved_value:
        # Industry benchmarks for Debtor Days
        if calculation.achieved_value <= 30:
            performance_level = 'Excellent'
            performance_color = 'success'
            performance_description = 'Outstanding cash collection efficiency'
        elif calculation.achieved_value <= 45:
            performance_level = 'Good'
            performance_color = 'primary'
            performance_description = 'Good payment collection performance'
        elif calculation.achieved_value <= 60:
            performance_level = 'Acceptable'
            performance_color = 'warning'
            performance_description = 'Acceptable collection timeframe'
        else:
            performance_level = 'Needs Improvement'
            performance_color = 'danger'
            performance_description = 'Collection period too long, requires attention'
    else:
        performance_level = 'No Data'
        performance_color = 'secondary'
        performance_description = 'Unable to assess performance'
    
    context = {
        'calculation': calculation,
        'performance_level': performance_level,
        'performance_color': performance_color,
        'performance_description': performance_description
    }
    return render(request, 'NAWEC_KPI/calculate_dd_detail.html', context)


@login_required
def calculate_dd_edit(request, calc_id):
    """Edit DD calculation"""
    calculation = get_object_or_404(CalculateDD, pk=calc_id)
    
    if request.method == 'POST':
        # Get form data
        trade_receivables = request.POST.get('trade_receivables')
        total_credit_sales = request.POST.get('total_credit_sales')
        year_id = request.POST.get('year')
        quarter_id = request.POST.get('quarter')
        
        try:
            # Update calculation fields
            calculation.trade_receivables = float(trade_receivables) if trade_receivables else 0
            calculation.total_credit_sales = float(total_credit_sales) if total_credit_sales else 0
            
            if year_id:
                calculation.year = YEAR.objects.get(id=year_id)
            if quarter_id:
                calculation.quarter = Quarter.objects.get(id=quarter_id)
            
            # Save will trigger automatic calculation via model save method
            calculation.save()
            
            messages.success(request, 'DD calculation updated successfully!')
            return redirect('NAWEC_KPI:calculate_dd_list')
            
        except Exception as e:
            messages.error(request, f'Error updating calculation: {str(e)}')
    
    # Get dropdown options
    years = YEAR.objects.all()
    quarters = Quarter.objects.all()
    
    context = {
        'calculation': calculation,
        'years': years,
        'quarters': quarters
    }
    return render(request, 'NAWEC_KPI/calculate_dd_edit.html', context)


@login_required
def calculate_dd_delete(request, calc_id):
    """Delete DD calculation"""
    calculation = get_object_or_404(CalculateDD, pk=calc_id)
    
    if request.method == 'POST':
        calculation.delete()
        messages.success(request, 'DD calculation deleted successfully!')
        return redirect('NAWEC_KPI:calculate_dd_list')
    
    context = {'calculation': calculation}
    return render(request, 'NAWEC_KPI/calculate_dd_delete.html', context)


# HTMX Cascading Dropdown Views
@login_required
def get_project_outcomes(request):
    """HTMX endpoint to get Project Outcomes filtered by project (PDO functionality removed)"""
    project_id = request.GET.get('project')
    if project_id:
        outcomes = ProjectOutCome.objects.filter(project_id=project_id).order_by('project_outcome')
    else:
        outcomes = ProjectOutCome.objects.none()
    
    return render(request, 'NAWEC_KPI/partials/outcome_options.html', {'outcomes': outcomes})

@login_required
def get_results_by_outcome(request):
    """HTMX endpoint to get Project Results filtered by Project Outcome"""
    outcome_id = request.GET.get('project_outcome')
    if outcome_id:
        results = ProjectResult.objects.filter(project_outcome_id=outcome_id).order_by('project_result')
    else:
        results = ProjectResult.objects.none()
    
    return render(request, 'NAWEC_KPI/partials/result_options.html', {'results': results})


# API Views for KPI Calculations
from django.views import View
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
import json

@method_decorator(csrf_exempt, name='dispatch')
class SaveKPICalculationView(View):
    """API endpoint to save KPI calculations from popup forms"""
    
    def get_quarter_object(self, quarter, kpi_type):
        """Helper method to map quarter string to Quarter object"""
        quarter_obj = None
        if quarter:
            try:
                # Map quarter string to actual database IDs
                quarter_mapping = {
                    '1': 10030,  # Quarter 1 (ID: 10030)
                    '2': 10031,  # Quarter 2 (ID: 10031)
                    '3': 10032,  # Quarter 3 (ID: 10032)
                    '4': 10033   # Quarter 4 (ID: 10033)
                }
                quarter_id = quarter_mapping.get(str(quarter))
                if quarter_id:
                    quarter_obj = Quarter.objects.get(id=quarter_id)
                    print(f"[DEBUG] {kpi_type} - Quarter mapped successfully: {quarter} -> {quarter_id}")
                else:
                    print(f"[DEBUG] {kpi_type} - Quarter mapping failed for: {quarter}")
            except Quarter.DoesNotExist:
                print(f"[DEBUG] {kpi_type} - Quarter ID {quarter_id} not found in database")
        return quarter_obj
    
    def post(self, request):
        try:
            data = json.loads(request.body)
            kpi_type = data.get('kpi_type')
            input_values = data.get('input_values', {})
            
            # Debug logging
            print(f"[DEBUG] SaveKPICalculationView - Raw data: {data}")
            print(f"[DEBUG] SaveKPICalculationView - KPI type: {kpi_type}")
            print(f"[DEBUG] SaveKPICalculationView - Input values: {input_values}")
            
            if not kpi_type:
                print(f"[DEBUG] SaveKPICalculationView - KPI type is missing or empty")
                return JsonResponse({'success': False, 'error': 'KPI type is required'})
            
            # Get the achieved value from input_values
            achieved_value = input_values.get('achieved_value')
            quarter = input_values.get('quarter')
            
            if not achieved_value:
                return JsonResponse({'success': False, 'error': 'Achieved value is required'})
            
            # Handle different KPI types
            if kpi_type == 'ROA':
                # ROA calculation
                net_profit_after_tax = input_values.get('net_income', 0)
                total_assets = input_values.get('total_assets', 0)
                compensation_amount = input_values.get('compensation_amount', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'ROA')
                
                calculation = CalculateROA(
                    net_profit_after_tax=net_profit_after_tax,
                    total_assets=total_assets,
                    compensation_amount=compensation_amount,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                print(f"[DEBUG] ROA calculation saved: ID={calculation.id}, ROA={calculation.achieved_value}%")
                
            elif kpi_type == 'NPM':
                # NPM calculation
                netprofit = input_values.get('netprofit', 0)
                total_revenues_turnover = input_values.get('total_revenues_turnover', 0)
                compensation_amount = input_values.get('compensation_amount', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'NPM')
                
                calculation = CalculateNPM(
                    netprofit=netprofit,
                    total_revenues_turnover=total_revenues_turnover,
                    compensation_amount=compensation_amount,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                print(f"[DEBUG] NPM calculation saved: ID={calculation.id}, NPM={calculation.achieved_value}%")
                
            elif kpi_type == 'DSCR':
                # DSCR calculation
                net_operating_income = input_values.get('net_operating_income', 0)
                total_debt_service = input_values.get('total_debt_service', 0)
                compensation_amount = input_values.get('compensation_amount', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'DSCR')
                
                calculation = CalculateDSCR(
                    net_operating_income=net_operating_income,
                    total_debt_service=total_debt_service,
                    compensation_amount=compensation_amount,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                print(f"[DEBUG] DSCR calculation saved: ID={calculation.id}, DSCR={calculation.achieved_value}")
                
            elif kpi_type == 'DER':
                # DER calculation
                total_debt = input_values.get('total_debt', 0)
                total_equity = input_values.get('total_equity', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'DER')
                
                calculation = CalculateDER(
                    total_debt=total_debt,
                    total_equity=total_equity,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'CR':
                # CR calculation
                current_assets = input_values.get('current_assets', 0)
                current_liabilities = input_values.get('current_liabilities', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'CR')
                
                calculation = CalculateCR(
                    current_assets=current_assets,
                    current_liabilities=current_liabilities,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'AO':
                # AO calculation
                audit_opinion = input_values.get('audit_opinion')
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'AO')
                
                calculation = CalculateAO(
                    audit_opinion=audit_opinion,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'PARI':
                # PARI calculation
                total_recommendations = input_values.get('total_recommendations', 0)
                total_implemented = input_values.get('total_implemented', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'PARI')
                
                calculation = CalculatePARI(
                    total_recommendations=total_recommendations,
                    total_implemented=total_implemented,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'TSQR':
                # TSQR calculation
                due_date = input_values.get('due_date', 0)
                actual_date = input_values.get('actual_date', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'TSQR')
                
                calculation = CalculateTSQR(
                    due_date=due_date,
                    actual_date=actual_date,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'DD':
                # DD calculation
                trade_receivables = input_values.get('trade_receivables', 0)
                total_credit_sales = input_values.get('total_credit_sales', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'DD')
                
                calculation = CalculateDD(
                    trade_receivables=trade_receivables,
                    total_credit_sales=total_credit_sales,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'WQCB':
                # WQCB calculation
                compliant_samples = input_values.get('compliant_samples', 0)
                total_samples = input_values.get('total_samples', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'WQCB')
                
                calculation = CalculateWQCB(
                    compliant_samples=compliant_samples,
                    total_samples=total_samples,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'WQCC':
                # WQCC calculation
                compliant_samples = input_values.get('compliant_samples', 0)
                total_samples = input_values.get('total_samples', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'WQCC')
                
                calculation = CalculateWQCC(
                    compliant_samples=compliant_samples,
                    total_samples=total_samples,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'EI':
                # EI (Energy Injection) calculation using CalculateMWh model
                # Handle dynamic energy sources from the new popup design
                total_energy = 0
                source_count = 0
                
                # Collect all energy source values from dynamic input
                for key, value in input_values.items():
                    if key.endswith('_energy') and key != 'achieved_value':
                        try:
                            energy_value = float(value)
                            if energy_value > 0:
                                total_energy += energy_value
                                source_count += 1
                        except (ValueError, TypeError):
                            continue
                
                # If no energy sources found, try the old format for backward compatibility
                if total_energy == 0:
                    solar_energy = input_values.get('solar_energy', 0)
                    wind_energy = input_values.get('wind_energy', 0)
                    thermal_energy = input_values.get('thermal_energy', 0)
                    other_energy = input_values.get('other_energy', 0)
                    hydro_energy = input_values.get('hydro_energy', 0)
                    nuclear_energy = input_values.get('nuclear_energy', 0)
                    
                    total_energy = solar_energy + wind_energy + thermal_energy + other_energy + hydro_energy + nuclear_energy
                    source_count = sum(1 for x in [solar_energy, wind_energy, thermal_energy, other_energy, hydro_energy, nuclear_energy] if x > 0)
                
                if total_energy <= 0:
                    return JsonResponse({'success': False, 'error': 'No valid energy sources provided'})
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'EI')
                
                calculation = CalculateMWh(
                    power_injected=total_energy,  # Use total energy as power injected
                    time_duration=1,  # Set as 1 hour for MW calculation
                    number_of_sources=max(source_count, 1),  # Number of actual energy sources
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'TTP':
                # TTP calculation
                on_time_payments = input_values.get('on_time_payments', 0)
                total_payments = input_values.get('total_payments', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'TTP')
                
                calculation = CalculateTTP(
                    on_time_payments=on_time_payments,
                    total_payments=total_payments,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'TPS':
                # TPS calculation
                on_time_payments = input_values.get('on_time_payments', 0)
                total_payments = input_values.get('total_payments', 0)
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'TPS')
                
                calculation = CalculateTPS(
                    on_time_payments=on_time_payments,
                    total_payments=total_payments,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'ATC':
                # ATC calculation
                billing_efficiency = input_values.get('billing_efficiency', 0)
                collection_efficiency = input_values.get('collection_efficiency', 0)
                
                # Calculate ATC&C: (1 - (billing_efficiency * collection_efficiency) / 10000) * 100
                # Convert percentages to decimals first, then calculate
                billing_decimal = billing_efficiency / 100
                collection_decimal = collection_efficiency / 100
                atc_value = (1 - (billing_decimal * collection_decimal)) * 100
                
                # Get Quarter object with proper mapping
                quarter_obj = self.get_quarter_object(quarter, 'ATC')
                
                calculation = CalculateATC(
                    billing_efficiency=billing_efficiency,
                    collection_efficiency=collection_efficiency,
                    achieved_value=atc_value,
                    quarter=quarter_obj,
                    loginUser=request.user
                )
                calculation.save()
                
            elif kpi_type == 'TMH':
                print(f"[DEBUG] TMH - Starting TMH calculation processing")
                try:
                    # TMH calculation with multi-session support
                    sessions = input_values.get('sessions', [])
                    total_sessions = input_values.get('total_sessions', 1)
                    achieved_value = input_values.get('achieved_value', 0)
                    
                    print(f"[DEBUG] TMH - Sessions data: {sessions}")
                    print(f"[DEBUG] TMH - Total sessions: {total_sessions}")
                    print(f"[DEBUG] TMH - Achieved value: {achieved_value}")
                    
                    # Get Quarter object with proper mapping
                    quarter_obj = self.get_quarter_object(quarter, 'TMH')
                    print(f"[DEBUG] TMH - Quarter object: {quarter_obj}")
                    
                    # For multi-session TMH, create a single record with aggregated data
                    if sessions and len(sessions) > 0:
                        print(f"[DEBUG] TMH - Processing multi-session data")
                        # Use first session for main record, combine data in title
                        first_session = sessions[0]
                        title = f"Multi-Session Training ({total_sessions} sessions)"
                        
                        # Aggregate session data for summary
                        total_participants = sum(session.get('number_of_participants', 0) for session in sessions)
                        total_days = sum(session.get('number_of_days', 0) for session in sessions)
                        avg_hours_per_day = sum(session.get('hours_per_day', 0) for session in sessions) / len(sessions)
                        
                        print(f"[DEBUG] TMH - Aggregated data: participants={total_participants}, days={total_days}, avg_hours={avg_hours_per_day}")
                        
                        # Handle date parsing safely
                        start_date = first_session.get('start_date')
                        end_date = sessions[-1].get('end_date')
                        
                        # Convert string dates to date objects if needed
                        from datetime import datetime
                        if isinstance(start_date, str):
                            try:
                                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                                print(f"[DEBUG] TMH - Parsed start_date: {start_date}")
                            except ValueError as e:
                                print(f"[DEBUG] TMH - Start date parsing error: {e}")
                                start_date = None
                        
                        if isinstance(end_date, str):
                            try:
                                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
                                print(f"[DEBUG] TMH - Parsed end_date: {end_date}")
                            except ValueError as e:
                                print(f"[DEBUG] TMH - End date parsing error: {e}")
                                end_date = None
                        
                        # Create CalculateTMH with aggregated multi-session data
                        calculation = CalculateTMH(
                            title=title,
                            start_date=start_date,
                            end_date=end_date,
                            hours_per_day=float(avg_hours_per_day),
                            number_of_participants=int(total_participants),
                            achieved_value=float(achieved_value),
                            quarter=quarter_obj,
                            loginUser=request.user
                        )
                        calculation.save()
                        print(f"[DEBUG] Multi-session TMH calculation saved: ID={calculation.id}, {achieved_value} man-hours for {total_sessions} sessions")
                    else:
                        print(f"[DEBUG] TMH - Processing single session fallback")
                        # Fallback for single session or missing data
                        calculation = CalculateTMH(
                            title="Training Session",
                            start_date=None,
                            end_date=None,
                            hours_per_day=8.0,
                            number_of_participants=1,
                            achieved_value=float(achieved_value),
                            quarter=quarter_obj,
                            loginUser=request.user
                        )
                        calculation.save()
                        print(f"[DEBUG] Single TMH calculation saved: ID={calculation.id}, {achieved_value} man-hours")
                except Exception as tmh_error:
                    print(f"[DEBUG] TMH - Exception occurred: {tmh_error}")
                    import traceback
                    print(f"[DEBUG] TMH - Traceback: {traceback.format_exc()}")
                    raise tmh_error
                
            else:
                print(f"[DEBUG] SaveKPICalculationView - Unhandled KPI type: '{kpi_type}'")
                return JsonResponse({'success': False, 'error': f'KPI type {kpi_type} not supported'})
            
            return JsonResponse({
                'success': True,
                'message': f'{kpi_type} calculation saved successfully',
                'achieved_value': achieved_value
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})


@method_decorator(csrf_exempt, name='dispatch')
class DeleteKPICalculationView(View):
    """API endpoint to delete KPI calculations"""
    
    def post(self, request):
        try:
            data = json.loads(request.body)
            kpi_type = data.get('kpi_type')
            calc_id = data.get('calc_id')
            
            if not kpi_type or not calc_id:
                return JsonResponse({'success': False, 'error': 'KPI type and calculation ID are required'})
            
            # Handle different KPI types
            if kpi_type == 'ROA':
                CalculateROA.objects.filter(id=calc_id).delete()
            elif kpi_type == 'DER':
                CalculateDER.objects.filter(id=calc_id).delete()
            elif kpi_type == 'CR':
                CalculateCR.objects.filter(id=calc_id).delete()
            elif kpi_type == 'AO':
                CalculateAO.objects.filter(id=calc_id).delete()
            elif kpi_type == 'EI':
                CalculateMWh.objects.filter(id=calc_id).delete()
            elif kpi_type == 'PARI':
                CalculatePARI.objects.filter(id=calc_id).delete()
            elif kpi_type == 'TSQR':
                CalculateTSQR.objects.filter(id=calc_id).delete()
            elif kpi_type == 'DD':
                CalculateDD.objects.filter(id=calc_id).delete()
            elif kpi_type == 'WQCB':
                CalculateWQCB.objects.filter(id=calc_id).delete()
            elif kpi_type == 'WQCC':
                CalculateWQCC.objects.filter(id=calc_id).delete()
            elif kpi_type == 'TTP':
                CalculateTTP.objects.filter(id=calc_id).delete()
            elif kpi_type == 'TPS':
                CalculateTPS.objects.filter(id=calc_id).delete()
            elif kpi_type == 'ATC':
                CalculateATC.objects.filter(id=calc_id).delete()
            elif kpi_type == 'TMH':
                CalculateTMH.objects.filter(id=calc_id).delete()
            elif kpi_type == 'IMPORTS':
                CalculateIMPORTS.objects.filter(id=calc_id).delete()
            elif kpi_type == 'IPP':
                CalculateIPP.objects.filter(id=calc_id).delete()
            else:
                return JsonResponse({'success': False, 'error': f'KPI type {kpi_type} not supported'})
            
            return JsonResponse({
                'success': True,
                'message': f'{kpi_type} calculation deleted successfully'
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})


# Imports (MW) Calculation Views
@login_required
def calculate_imports_list(request):
    """List view for Imports (MW) calculations with search and filtering"""
    search_query = request.GET.get('search', '')
    year_filter = request.GET.get('year', '')
    quarter_filter = request.GET.get('quarter', '')
    
    # Base queryset
    calculations = CalculateIMPORTS.objects.all().order_by('-date_created')
    
    # Apply filters
    if search_query:
        calculations = calculations.filter(
            Q(End_Target_Value__icontains=search_query) |
            Q(add_value__icontains=search_query) |
            Q(achieved_value__icontains=search_query) |
            Q(loginUser__username__icontains=search_query)
        )
    
    if year_filter:
        calculations = calculations.filter(year_id=year_filter)
    
    if quarter_filter:
        calculations = calculations.filter(quarter_id=quarter_filter)
    
    # Pagination
    paginator = Paginator(calculations, 25)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Get filter options
    years = YEAR.objects.all().order_by('-profile_year')
    quarters = Quarter.objects.filter(id__in=[10030, 10031, 10032, 10033]).order_by('id')
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'year_filter': year_filter,
        'quarter_filter': quarter_filter,
        'years': years,
        'quarters': quarters,
        'total_count': calculations.count(),
    }
    
    return render(request, 'NAWEC_KPI/calculate_imports_list.html', context)


@login_required
def calculate_imports_detail(request, pk):
    """Detail view for specific Imports (MW) calculation"""
    try:
        calculation = CalculateIMPORTS.objects.get(pk=pk)
    except CalculateIMPORTS.DoesNotExist:
        messages.error(request, f'Imports (MW) calculation with ID {pk} does not exist.')
        return redirect('NAWEC_KPI:calculate_imports_list')
    
    context = {
        'calculation': calculation,
    }
    
    return render(request, 'NAWEC_KPI/calculate_imports_detail.html', context)


@login_required
def calculate_imports_edit(request, pk):
    """Edit view for Imports (MW) calculation"""
    try:
        calculation = CalculateIMPORTS.objects.get(pk=pk)
    except CalculateIMPORTS.DoesNotExist:
        messages.error(request, f'Imports (MW) calculation with ID {pk} does not exist.')
        return redirect('NAWEC_KPI:calculate_imports_list')
    
    if request.method == 'POST':
        # Update calculation fields
        try:
            end_target = float(request.POST.get('End_Target_Value', 0))
            add_value = float(request.POST.get('add_value', 0))
            
            calculation.End_Target_Value = end_target
            calculation.add_value = add_value
            
            # Recalculate achieved value
            calculation.achieved_value = end_target * add_value
            
            calculation.save()
            
            messages.success(request, 'Imports (MW) calculation updated successfully!')
            return redirect('NAWEC_KPI:calculate_imports_detail', pk=pk)
            
        except (ValueError, TypeError) as e:
            messages.error(request, f'Invalid input values: {str(e)}')
    
    context = {
        'calculation': calculation,
    }
    
    return render(request, 'NAWEC_KPI/calculate_imports_edit.html', context)


@login_required
def calculate_imports_delete(request, pk):
    """Delete view for Imports (MW) calculation"""
    try:
        calculation = CalculateIMPORTS.objects.get(pk=pk)
    except CalculateIMPORTS.DoesNotExist:
        messages.error(request, f'Imports (MW) calculation with ID {pk} does not exist.')
        return redirect('NAWEC_KPI:calculate_imports_list')
    
    if request.method == 'POST':
        calculation.delete()
        messages.success(request, 'Imports (MW) calculation deleted successfully!')
        return redirect('NAWEC_KPI:calculate_imports_list')
    
    context = {
        'calculation': calculation,
    }
    
    return render(request, 'NAWEC_KPI/calculate_imports_delete.html', context)


# Independent Power Plants (MW) CRUD Views
@login_required
def calculate_ipp_list(request):
    """List view for Independent Power Plants (MW) calculations"""
    calculations = CalculateIPP.objects.all().order_by('-date_created')
    
    # Pagination
    paginator = Paginator(calculations, 10)  # 10 calculations per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'calculations': page_obj,
        'page_obj': page_obj,
        'title': 'Independent Power Plants (MW) Calculations'
    }
    
    return render(request, 'NAWEC_KPI/calculate_ipp_list.html', context)


@login_required
def calculate_ipp_detail(request, pk):
    """Detail view for Independent Power Plants (MW) calculation"""
    try:
        calculation = CalculateIPP.objects.get(pk=pk)
    except CalculateIPP.DoesNotExist:
        messages.error(request, f'Independent Power Plants (MW) calculation with ID {pk} does not exist.')
        return redirect('NAWEC_KPI:calculate_ipp_list')
    
    context = {
        'calculation': calculation,
    }
    
    return render(request, 'NAWEC_KPI/calculate_ipp_detail.html', context)


@login_required
def calculate_ipp_edit(request, pk):
    """Edit view for Independent Power Plants (MW) calculation"""
    try:
        calculation = CalculateIPP.objects.get(pk=pk)
    except CalculateIPP.DoesNotExist:
        messages.error(request, f'Independent Power Plants (MW) calculation with ID {pk} does not exist.')
        return redirect('NAWEC_KPI:calculate_ipp_list')
    
    if request.method == 'POST':
        try:
            # Update calculation values
            calculation.End_Target_Value = float(request.POST.get('End_Target_Value', 0))
            calculation.add_value = float(request.POST.get('add_value', 0))
            
            # Recalculate the achieved value
            calculation.achieved_value = calculation.End_Target_Value * calculation.add_value
            
            calculation.save()
            messages.success(request, 'Independent Power Plants (MW) calculation updated successfully!')
            return redirect('NAWEC_KPI:calculate_ipp_detail', pk=pk)
        except (ValueError, TypeError) as e:
            messages.error(request, f'Invalid input values: {str(e)}')
    
    context = {
        'calculation': calculation,
    }
    
    return render(request, 'NAWEC_KPI/calculate_ipp_edit.html', context)


@login_required
def calculate_ipp_delete(request, pk):
    """Delete view for Independent Power Plants (MW) calculation"""
    try:
        calculation = CalculateIPP.objects.get(pk=pk)
    except CalculateIPP.DoesNotExist:
        messages.error(request, f'Independent Power Plants (MW) calculation with ID {pk} does not exist.')
        return redirect('NAWEC_KPI:calculate_ipp_list')
    
    if request.method == 'POST':
        calculation.delete()
        messages.success(request, 'Independent Power Plants (MW) calculation deleted successfully!')
        return redirect('NAWEC_KPI:calculate_ipp_list')
    
    context = {
        'calculation': calculation,
    }
    
    return render(request, 'NAWEC_KPI/calculate_ipp_delete.html', context)



